"""
generate_master_curriculum.py
Biên soạn tài liệu MASTER GIÁO TRÌNH & BÀI TẬP PYTHON NÂNG CAO
Đơn vị: TIN HỌC SAO VIỆT THỦ ĐỨC

CẤU TRÚC ĐẶC BIỆT:
1. CHƯƠNG 5 (ĐỒ HỌA TURTLE) ĐƯỢC ĐƯA LÊN ĐẦU TIÊN THEO YÊU CẦU.
2. TIẾP THEO LÀ BÀI 1 ĐẾN BÀI 11 (Chuỗi, List & Dict, Hàm, Thư viện).
3. MỖI CHƯƠNG BỔ SUNG THÊM ÍT NHẤT 2 BÀI TẬP MỞ RỘNG DẠNG MỚI (KHÔNG TRÙNG LẶP).
4. 6 DỰ ÁN BÀI TẬP NÂNG CAO TỔNG HỢP LIÊN MÔN.
5. KHO 10 BÀI TOÁN THỰC HÀNH TỰ LUẬN ĐỀ THI.
6. 100% CODE KHÔNG CÓ CHÚ THÍCH (#), GIẢI THÍCH CHI TIẾT & SUY LUẬN ĐẶT SAU OUTPUT.
7. ĐỊNH DẠNG IN ẤN TRẮNG ĐEN (HIGH CONTRAST B&W) CHUẨN A4 SẮC NÉT.
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx2pdf import convert

sys.stdout.reconfigure(encoding='utf-8')

CUR_DIR = os.path.dirname(os.path.abspath(__file__))
PKG_DIR = os.path.dirname(CUR_DIR)
TAI_LIEU_DIR = os.path.join(PKG_DIR, "Tai_Lieu_In_An")
os.makedirs(TAI_LIEU_DIR, exist_ok=True)

DOCX_PATH = os.path.join(TAI_LIEU_DIR, "Giao_Trinh_Bai_Tap_Python_Nang_Cao_Sao_Viet.docx")
PDF_PATH  = os.path.join(TAI_LIEU_DIR, "Giao_Trinh_Bai_Tap_Python_Nang_Cao_Sao_Viet.pdf")

BLACK     = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(50, 50, 50)
WHITE     = RGBColor(255, 255, 255)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_border(cell, top="single", bottom="single", left="single", right="single", sz="12", color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tcBorders>''')
    tcPr.append(tcBorders)

def add_header_banner(doc, title_text, subtitle_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "000000")
    set_cell_margins(cell, top=160, bottom=160, left=180, right=180)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(3)
    run1 = p1.add_run("TIN HỌC SAO VIỆT THỦ ĐỨC — TÀI LIỆU THỰC HÀNH LẬP TRÌNH PYTHON NÂNG CAO")
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    run2 = p2.add_run(title_text)
    run2.font.size = Pt(14.5)
    run2.font.bold = True
    run2.font.color.rgb = WHITE

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    run3 = p3.add_run(subtitle_text)
    run3.font.size = Pt(9.5)
    run3.font.italic = True
    run3.font.color.rgb = WHITE

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_chapter_divider(doc, chapter_title, sub_desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_title = p.add_run(f"═════ {chapter_title.upper()} ═════\n")
    r_title.font.size = Pt(12.5)
    r_title.font.bold = True
    r_title.font.color.rgb = BLACK
    
    r_desc = p.add_run(sub_desc)
    r_desc.font.size = Pt(9.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = DARK_GRAY

def add_lesson_box_bw(doc, badge, title, goal, metaphor, code_text, output_text, explanations, tip_text=None):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(3)
    p_title.paragraph_format.keep_with_next = True
    
    r_badge = p_title.add_run(f"▶ {badge}: ")
    r_badge.font.bold = True
    r_badge.font.size = Pt(11)
    r_badge.font.color.rgb = BLACK
    
    r_title = p_title.add_run(f"{title}")
    r_title.font.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = BLACK

    tbl_intro = doc.add_table(rows=1, cols=1)
    tbl_intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_intro = tbl_intro.cell(0, 0)
    set_cell_background(c_intro, "FFFFFF")
    set_cell_border(c_intro, top="single", bottom="single", left="single", right="single", sz="8", color="000000")
    set_cell_margins(c_intro, top=60, bottom=60, left=100, right=100)

    p_goal = c_intro.paragraphs[0]
    p_goal.paragraph_format.space_after = Pt(2)
    r_g_lbl = p_goal.add_run("• Mục tiêu học tập: ")
    r_g_lbl.font.bold = True
    p_goal.add_run(goal)

    if metaphor:
        p_meta = c_intro.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(0)
        r_m_lbl = p_meta.add_run("• Trực quan dễ nhớ: ")
        r_m_lbl.font.bold = True
        r_meta = p_meta.add_run(metaphor)
        r_meta.font.italic = True

    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_before = Pt(5)
    p_lbl.paragraph_format.space_after = Pt(2)
    p_lbl.paragraph_format.keep_with_next = True
    r_l = p_lbl.add_run("⌨️ MÃ NGUỒN PYTHON (GÕ VÀO MÁY):")
    r_l.font.bold = True
    r_l.font.size = Pt(9.5)

    tbl_code = doc.add_table(rows=1, cols=1)
    tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_code = tbl_code.cell(0, 0)
    set_cell_background(c_code, "FFFFFF")
    set_cell_border(c_code, top="single", bottom="single", left="single", right="single", sz="12", color="000000")
    set_cell_margins(c_code, top=80, bottom=80, left=120, right=120)

    p_code = c_code.paragraphs[0]
    p_code.paragraph_format.space_after = Pt(0)
    p_code.paragraph_format.line_spacing = 1.15
    run_code = p_code.add_run(code_text.strip())
    run_code.font.name = "Consolas"
    run_code.font.size = Pt(10)
    run_code.font.bold = False
    run_code.font.color.rgb = BLACK

    tbl_out = doc.add_table(rows=1, cols=1)
    tbl_out.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_out = tbl_out.cell(0, 0)
    set_cell_background(c_out, "F4F4F4")
    set_cell_border(c_out, top="dashed", bottom="dashed", left="dashed", right="dashed", sz="8", color="000000")
    set_cell_margins(c_out, top=60, bottom=60, left=100, right=100)

    p_out = c_out.paragraphs[0]
    p_out.paragraph_format.space_after = Pt(0)
    r_out_h = p_out.add_run("🖥️ KẾT QUẢ IN RA MÀN HÌNH (OUTPUT):\n")
    r_out_h.font.bold = True
    r_out_h.font.size = Pt(8.5)
    r_out_h.font.color.rgb = BLACK

    r_out_body = p_out.add_run(output_text.strip())
    r_out_body.font.name = "Consolas"
    r_out_body.font.size = Pt(9.5)
    r_out_body.font.color.rgb = BLACK

    if explanations:
        p_exp_h = doc.add_paragraph()
        p_exp_h.paragraph_format.space_before = Pt(4)
        p_exp_h.paragraph_format.space_after = Pt(2)
        p_exp_h.paragraph_format.keep_with_next = True
        r_eh = p_exp_h.add_run("📖 GIẢI THÍCH CHI TIẾT & SUY LUẬN LOGIC:")
        r_eh.font.bold = True
        r_eh.font.size = Pt(9.5)

        tbl_exp = doc.add_table(rows=1, cols=1)
        tbl_exp.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_exp = tbl_exp.cell(0, 0)
        set_cell_background(c_exp, "FFFFFF")
        set_cell_border(c_exp, top="single", bottom="single", left="single", right="single", sz="6", color="666666")
        set_cell_margins(c_exp, top=50, bottom=50, left=100, right=100)

        for idx, exp_line in enumerate(explanations):
            p_e = c_exp.paragraphs[0] if idx == 0 else c_exp.add_paragraph()
            p_e.paragraph_format.space_after = Pt(2)
            r_eb = p_e.add_run(f"• {exp_line}")
            r_eb.font.size = Pt(9)

    if tip_text:
        p_tip = doc.add_paragraph()
        p_tip.paragraph_format.space_before = Pt(3)
        p_tip.paragraph_format.space_after = Pt(8)
        r_t_h = p_tip.add_run("✍️ Lưu ý quan trọng cho học viên: ")
        r_t_h.font.bold = True
        r_t_h.font.size = Pt(9)
        r_t_b = p_tip.add_run(tip_text)
        r_t_b.font.size = Pt(9)
        r_t_b.font.italic = True

def add_kpi_table(doc, kpi_list):
    add_chapter_divider(doc, "BẢNG ĐÁNH GIÁ RÈN LUYỆN TOÀN KHÓA HỌC", "Tích chọn và chấm điểm sau khi hoàn thành từng bài trên máy")
    tbl_kpi = doc.add_table(rows=len(kpi_list) + 1, cols=3)
    tbl_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_headers = ["Nội Dung Bài Học & Dự Án Thực Hành", "Học Viên Tự Đánh Giá", "Thầy/Cô Nhận Xét"]
    for idx, h in enumerate(kpi_headers):
        c = tbl_kpi.cell(0, idx)
        set_cell_background(c, "000000")
        set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE

    for r_idx, item in enumerate(kpi_list):
        c0 = tbl_kpi.cell(r_idx + 1, 0)
        c1 = tbl_kpi.cell(r_idx + 1, 1)
        c2 = tbl_kpi.cell(r_idx + 1, 2)
        c0.width = Inches(4.2)
        c1.width = Inches(1.4)
        c2.width = Inches(1.6)
        set_cell_background(c0, "FFFFFF")
        set_cell_background(c1, "FFFFFF")
        set_cell_background(c2, "FFFFFF")
        set_cell_border(c0, sz="6", color="000000")
        set_cell_border(c1, sz="6", color="000000")
        set_cell_border(c2, sz="6", color="000000")
        set_cell_margins(c0, top=40, bottom=40, left=60, right=60)
        set_cell_margins(c1, top=40, bottom=40, left=60, right=60)
        set_cell_margins(c2, top=40, bottom=40, left=60, right=60)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p0.add_run(item)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run("[  ] Hoàn thành")

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run("⭐⭐⭐")


def build_master_doc():
    print("🚀 Bắt đầu tạo tài liệu MASTER CURRICULUM theo thứ tự chuẩn mới...")
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    font.color.rgb = BLACK

    add_header_banner(
        doc,
        "GIÁO TRÌNH & BÀI TẬP LẬP TRÌNH PYTHON NÂNG CAO",
        "Bộ Tài Liệu Thực Hành Chuẩn — Tin Học Sao Việt Thủ Đức — In Ấn Sắc Nét"
    )

    # ==========================================================================
    # PHẦN 1: CHƯƠNG 5 — THƯ VIỆN ĐỒ HỌA TURTLE (ĐƯA LÊN ĐẦU TIÊN)
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 1: CHƯƠNG 5 — THƯ VIỆN ĐỒ HỌA TURTLE (RÙA VẼ TRANH)",
        "Lý thuyết trọng tâm, các câu lệnh điều khiển rùa, 4 bài tập giáo trình gốc và 2 bài tập mở rộng mới."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 12",
        title="THƯ VIỆN TURTLE — KHỞI ĐỘNG THƯ VIỆN ĐỒ HỌA",
        goal="Khởi tạo đối tượng Turtle, đặt hình dạng chú rùa và giữ màn hình đồ họa luôn mở.",
        metaphor="Chú robot rùa thông minh cầm cây bút lông nhiều màu sẵn sàng vẽ tranh.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.shape("turtle")
but_ve.speed(3)

but_ve.forward(100)

turtle.done()""",
        output_text="""(Cửa sổ đồ họa mở ra: Chú rùa tiến thẳng 100 bước để lại một nét vẽ màu đen)""",
        explanations=[
            "import turtle: Nạp thư viện đồ họa rùa Turtle có sẵn trong Python.",
            "but_ve = turtle.Turtle(): Khởi tạo đối tượng bút vẽ điều khiển chú rùa.",
            "but_ve.shape('turtle'): Đổi con trỏ thành hình chú rùa dễ thương.",
            "turtle.done(): Giữ cửa sổ vẽ luôn mở để học viên quan sát tác phẩm."
        ],
        tip_text="Luôn luôn đặt lệnh turtle.done() ở dòng cuối cùng của chương trình Turtle."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 13",
        title="CÁC CÂU LỆNH ĐIỀU KHIỂN RÙA VẼ CƠ BẢN",
        goal="Làm chủ nhóm lệnh di chuyển forward, backward, left, right, circle; độ dày pensize và màu nét color.",
        metaphor="Bé điều khiển ô tô đồ chơi: tiến tới, lùi lại, rẽ góc và bấm nút phun sơn.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(3)
but_ve.color("blue")

but_ve.forward(100)
but_ve.right(90)
but_ve.forward(50)
but_ve.circle(40)

turtle.done()""",
        output_text="""(Chú rùa vẽ đường thẳng 100 bước màu xanh dương, rẽ phải 90 độ vẽ 50 bước và vẽ 1 hình tròn bán kính 40)""",
        explanations=[
            "but_ve.pensize(3): Đặt độ dày của nét vẽ bằng 3 điểm ảnh.",
            "but_ve.color('blue'): Đổi màu nét vẽ sang màu xanh dương.",
            "but_ve.forward(100): Tiến thẳng 100 bước.",
            "but_ve.right(90): Xoay góc vuông sang phải 90 độ.",
            "but_ve.circle(40): Vẽ hình tròn bán kính 40 bước."
        ],
        tip_text="but_ve.penup() nhấc bút không vẽ nét, but_ve.pendown() hạ bút để tiếp tục vẽ."
    )

    # 4 bài tập gốc Chương 5
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 1 (GỐC)",
        title="VẼ HÌNH VUÔNG NHIỀU MÀU (4 CẠNH 4 MÀU RỰC RỠ)",
        goal="Vẽ hình vuông cạnh 100 bước, mỗi cạnh một màu: đỏ, xanh lá, xanh dương, vàng.",
        metaphor="Khung tranh 4 màu sắc cầu vồng do chính tay bé lập trình vẽ.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(4)
but_ve.speed(3)

cac_mau = ["red", "green", "blue", "yellow"]

for mau in cac_mau:
    but_ve.color(mau)
    but_ve.forward(100)
    but_ve.right(90)

turtle.done()""",
        output_text="""(Màn hình hiển thị hình vuông cạnh 100 bước với 4 cạnh mang 4 màu Đỏ, Xanh lá, Xanh dương và Vàng)""",
        explanations=[
            "cac_mau = ['red', 'green', 'blue', 'yellow']: Danh sách 4 màu theo đề bài.",
            "Vòng lặp for mau in cac_mau đổi màu bút ở mỗi cạnh rồi tiến 100 bước và xoay phải 90 độ."
        ],
        tip_text="Góc quay ngoài của hình vuông là 90 độ (360 / 4 cạnh = 90 độ)."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 2 (GỐC)",
        title="VẼ HÌNH CHỮ NHẬT NÉT ĐẬM MÀU TÍM",
        goal="Vẽ hình chữ nhật dài 150 bước, rộng 80 bước, nét vẽ dày 5 điểm ảnh, màu tím (purple).",
        metaphor="Bé vẽ cánh cửa lớn của ngôi nhà cổ tích bằng nét bút lông màu tím cực đậm.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(5)
but_ve.color("purple")
but_ve.speed(3)

for i in range(2):
    but_ve.forward(150)
    but_ve.right(90)
    but_ve.forward(80)
    but_ve.right(90)

turtle.done()""",
        output_text="""(Màn hình hiển thị hình chữ nhật chiều dài 150, chiều rộng 80, nét vẽ dày 5 và có màu tím nổi bật)""",
        explanations=[
            "but_ve.pensize(5): Nét vẽ dày 5 điểm ảnh giúp đường nét to rõ khi in ấn.",
            "but_ve.color('purple'): Chọn màu mực vẽ tím.",
            "Vòng lặp for i in range(2) vẽ 2 lần cặp cạnh dài 150 và cạnh rộng 80."
        ],
        tip_text="Hình chữ nhật có 2 cặp cạnh đối diện bằng nhau."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 3 (GỐC)",
        title="VẼ HAI HÌNH VUÔNG XA NHAU (KỸ THUẬT PENUP & PENDOWN)",
        goal="Vẽ hình vuông 1 (cạnh 80), nhấc bút di chuyển sang phải 200 bước mà không vẽ nét, vẽ hình vuông 2 (cạnh 80, màu cam).",
        metaphor="Chú rùa vẽ xong một bức tranh, sau đó nhảy cóc sang vị trí bên cạnh để vẽ tiếp hình thứ hai.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(3)
but_ve.speed(3)

but_ve.color("blue")
for i in range(4):
    but_ve.forward(80)
    but_ve.right(90)

but_ve.penup()
but_ve.forward(200)
but_ve.pendown()

but_ve.color("orange")
for i in range(4):
    but_ve.forward(80)
    but_ve.right(90)

turtle.done()""",
        output_text="""(Màn hình hiển thị 2 hình vuông riêng biệt cạnh 80 bước cách nhau 200 bước: hình 1 màu xanh dương, hình 2 màu cam)""",
        explanations=[
            "Vẽ hình vuông thứ nhất màu xanh dương cạnh 80 bước.",
            "but_ve.penup(): Nhấc bút lên, di chuyển 200 bước mà không để lại nét mực thừa.",
            "but_ve.pendown(): Hạ bút xuống và vẽ hình vuông thứ hai màu cam."
        ],
        tip_text="Kỹ thuật penup() và pendown() là kỹ năng bắt buộc khi vẽ các họa tiết tách rời."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 4 (GỐC)",
        title="VẼ TAM GIÁC ĐỀU & ĐỔI MÀU NỀN MÀN HÌNH (BGCOLOR & TÔ MÀU)",
        goal="Đổi màu nền màn hình thành 'lightblue', vẽ hình tam giác đều cạnh 120 bước, nét dày 3, viền và nền tô màu đỏ.",
        metaphor="Lá cờ tam giác màu đỏ rực rỡ nổi bật trên nền trời xanh biếc mùa hè.",
        code_text="""import turtle

man_hinh = turtle.Screen()
man_hinh.bgcolor("lightblue")

but_ve = turtle.Turtle()
but_ve.pensize(3)
but_ve.color("red", "red")
but_ve.speed(3)

but_ve.begin_fill()
for i in range(3):
    but_ve.forward(120)
    but_ve.left(120)
but_ve.end_fill()

turtle.done()""",
        output_text="""(Cửa sổ nền xanh nhạt 'lightblue' hiển thị hình tam giác đều màu đỏ rực rỡ, cạnh 120 bước, viền dày 3)""",
        explanations=[
            "man_hinh.bgcolor('lightblue'): Đổi màu nền toàn bộ cửa sổ đồ họa.",
            "but_ve.color('red', 'red'): Thiết lập màu viền và màu tô đều là đỏ.",
            "but_ve.begin_fill() và but_ve.end_fill(): Bắt đầu và kết thúc việc tô màu kín.",
            "Góc quay ngoài của tam giác đều là 120 độ (360 / 3 = 120 độ)."
        ],
        tip_text="Tổng 3 góc ngoài của mọi hình đa giác đều luôn luôn bằng 360 độ."
    )

    # 2 bài tập mở rộng Chương 5
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 5 (MỞ RỘNG MỚI)",
        title="VẼ NGÔI SAO 5 CÁNH VÀNG VIỀN ĐỎ RỰC RỠ",
        goal="Vẽ ngôi sao 5 cánh có độ dài cạnh 150 bước, góc xoay 144 độ, viền đỏ và tô màu vàng bên trong.",
        metaphor="Ngôi sao lấp lánh trên bầu trời đêm hoặc ngôi sao vàng trên quốc kỳ.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(3)
but_ve.speed(3)
but_ve.color("red", "yellow")

but_ve.begin_fill()
for i in range(5):
    but_ve.forward(150)
    but_ve.right(144)
but_ve.end_fill()

turtle.done()""",
        output_text="""(Màn hình hiển thị ngôi sao 5 cánh rực rỡ với đường viền màu đỏ và phần ruột được tô vàng óng ánh)""",
        explanations=[
            "Góc quay đặc trưng của ngôi sao 5 cánh chuẩn là 144 độ (right(144)).",
            "Vòng lặp for chạy đúng 5 lần tương ứng với 5 đỉnh của ngôi sao.",
            "but_ve.begin_fill() và but_ve.end_fill() tự động tô kín màu vàng cho toàn bộ thân ngôi sao."
        ],
        tip_text="Ngôi sao 5 cánh được tạo nên bằng cách vẽ đan chéo các đoạn thẳng xoay góc 144 độ."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 5 — BÀI TẬP 6 (MỞ RỘNG MỚI)",
        title="VẼ HỌA TIẾT XOẮN ỐC NGHỆ THUẬT TĂNG DẦN BÁN KÍNH",
        goal="Sử dụng vòng lặp for tăng dần kích thước bước đi để tạo nên hình xoắn ốc đối xứng tuyệt đẹp.",
        metaphor="Vỏ ốc biển kỳ diệu cuộn tròn mở rộng dần từ tâm ra ngoài.",
        code_text="""import turtle

but_ve = turtle.Turtle()
but_ve.pensize(2)
but_ve.speed(0)
but_ve.color("darkblue")

for i in range(1, 80):
    but_ve.forward(i * 3)
    but_ve.right(91)

turtle.done()""",
        output_text="""(Cửa sổ Turtle vẽ một bức tranh xoắn ốc hình học nghệ thuật xoay đều từ tâm lan tỏa ra ngoài)""",
        explanations=[
            "Biến đếm i tăng dần từ 1 đến 79 giúp độ dài bước đi (i * 3) dài ra sau mỗi góc rẽ.",
            "Góc xoay 91 độ (lớn hơn góc vuông 90 độ đúng 1 độ) tạo nên hiệu ứng xoắn lệch tâm nghệ thuật.",
            "Tốc độ speed(0) giúp máy tính vẽ nhanh như chớp."
        ],
        tip_text="Thay đổi góc xoay thành 61 độ hoặc 121 độ sẽ tạo ra các hình xoắn tam giác hoặc lục giác rất lạ mắt."
    )


    # ==========================================================================
    # PHẦN 2: CHƯƠNG 1 — CÁC THAO TÁC VỚI KIỂU DỮ LIỆU STRING
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 2: CHƯƠNG 1 — CÁC THAO TÁC VỚI KIỂU DỮ LIỆU STRING (CHUỖI)",
        "Học phần Bài 1, Bài 2, Bài 3, 3 bài tập giáo trình gốc và 2 bài tập mở rộng mới."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 1",
        title="CƠ BẢN VỀ CHUỖI (STRING) & TRUY CẬP KÝ TỰ QUA VỊ TRÍ",
        goal="Tạo chuỗi, sử dụng ký tự đặc biệt '\\n', '\\t' và cắt chuỗi theo chỉ số index [0], [-1], [start:end].",
        metaphor="Chuỗi như đoàn tàu hỏa gồm nhiều toa, mỗi toa mang 1 chữ cái được đánh số vị trí từ 0.",
        code_text="""s = "Tin Hoc Sao Viet"
print(s)
print(s[0])
print(s[-1])
print(s[0:7])
print("Dong 1\\nDong 2\\tTab")""",
        output_text="""Tin Hoc Sao Viet
T
t
Tin Hoc
Dong 1
Dong 2\tTab""",
        explanations=[
            "s[0]: Lấy ký tự đầu tiên tại vị trí 0 (chữ 'T').",
            "s[-1]: Lấy ký tự cuối cùng của chuỗi (chữ 't').",
            "s[0:7]: Cắt chuỗi từ vị trí 0 đến trước vị trí 7 thu được 'Tin Hoc'.",
            "\\n là ký tự xuống dòng mới, \\t là ký tự thụt lề Tab."
        ],
        tip_text="Chuỗi trong Python là bất biến (immutable), không thể sửa trực tiếp từng ký tự."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 2",
        title="CÁC PHƯƠNG THỨC XỬ LÝ CHUỖI THƯỜNG DÙNG",
        goal="Làm chủ các phương thức upper(), lower(), title(), strip(), len(), find(), replace() và kiểm tra chuỗi.",
        metaphor="Các nút bấm chức năng tự động: in hoa toàn bộ, xóa khoảng trắng thừa, tìm kiếm từ.",
        code_text="""s = "  tin hoc sao viet  "
print(s.upper())
print(s.lower())
print(s.title())
print(s.strip())
print(len(s))
print(s.find("sao"))
print(s.replace("viet", "nam"))""",
        output_text="""  TIN HOC SAO VIET  
  tin hoc sao viet  
  Tin Hoc Sao Viet  
tin hoc sao viet
20
10
  tin hoc sao nam  """,
        explanations=[
            "s.upper(): Chuyển tất cả chữ thành chữ hoa in đậm nét.",
            "s.lower(): Chuyển tất cả chữ thành chữ thường.",
            "s.title(): Tự động viết hoa chữ cái đầu tiên của từng từ.",
            "s.strip(): Cắt bỏ khoảng trắng dư thừa ở đầu và cuối chuỗi.",
            "len(s): Đếm tổng số lượng ký tự có trong chuỗi.",
            "s.find('sao'): Trả về vị trí xuất hiện đầu tiên của chữ 'sao'.",
            "s.replace('viet', 'nam'): Thay thế chữ 'viet' bằng chữ 'nam'."
        ],
        tip_text="Hàm s.isdigit() kiểm tra toàn bộ là số, s.isalpha() kiểm tra toàn bộ là chữ cái."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 3",
        title="ĐỊNH DẠNG CHUỖI (STRING FORMATTING) VỚI F-STRING",
        goal="Truyền biến vào văn bản nhanh chóng bằng phương thức .format() và f-string chuyên nghiệp.",
        metaphor="Mẫu giấy khen in sẵn, em chỉ việc điền Tên và Điểm số vào đúng chỗ trống.",
        code_text="""ten = "Bao Nam"
tuoi = 10
diem = 9.8765

print("Ban {} nam nay {} tuoi".format(ten, tuoi))
print(f"Ban {ten} nam nay {tuoi} tuoi")
print(f"Diem so lam tron 2 chu so: {diem:.2f}")""",
        output_text="""Ban Bao Nam nam nay 10 tuoi
Ban Bao Nam nam nay 10 tuoi
Diem so lam tron 2 chu so: 9.88""",
        explanations=[
            ".format(ten, tuoi): Điền tuần tự giá trị của biến ten và tuoi vào dấu ngoặc {}.",
            "f\"Ban {ten}...\": Cú pháp f-string hiện đại, đặt biến trực tiếp trong ngoặc nhọn.",
            "{diem:.2f}: Định dạng số thực làm tròn đúng 2 chữ số thập phân sau dấu phẩy."
        ],
        tip_text="Khi dùng f-string bắt buộc phải có chữ 'f' hoặc 'F' đứng ngay trước dấu ngoặc kép."
    )

    # 3 bài tập gốc Chương 1
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 1 — BÀI TẬP 1 (GỐC)",
        title="KIỂM TRA TÊN NGƯỜI DÙNG CÓ SỐ VÀ KHOẢNG TRẮNG",
        goal="Nhập tên từ bàn phím, kiểm tra tên có chứa chữ số và khoảng trắng để in thông báo phân loại.",
        metaphor="Cổng soát vé tự động quét xem họ tên học sinh có bị gõ nhầm số hay không.",
        code_text="""ten = input("Nhap ho ten: ")

co_so = False
for ch in ten:
    if ch.isdigit():
        co_so = True
        break

co_khoang_trang = " " in ten

if co_so and not co_khoang_trang:
    print("Ten nguoi dung co so va khong co khoang trang")
elif not co_so and co_khoang_trang:
    print("Ten nguoi dung khong co so va co khoang trang")
elif co_so and co_khoang_trang:
    print("Ten nguoi dung co ca so va khoang trang")
else:
    print("Ten nguoi dung chi gom chu cai, khong co so va khong co khoang trang")""",
        output_text="""Nhap ho ten: NguyenVanAn1
Ten nguoi dung co so va khong co khoang trang

Nhap ho ten: Nguyen Van An
Ten nguoi dung khong co so va co khoang trang""",
        explanations=[
            "Vòng lặp for ch in ten duyệt từng ký tự, dùng ch.isdigit() phát hiện chữ số.",
            "Cú pháp ' ' in ten kiểm tra xem có ký tự khoảng cách hay không.",
            "Cấu trúc if-elif-else phân loại 4 trường hợp dữ liệu."
        ],
        tip_text="Toán tử 'in' trong Python giúp kiểm tra sự tồn tại của chuỗi con cực nhanh."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 1 — BÀI TẬP 2 (GỐC)",
        title="CHUẨN HÓA HỌ TÊN HỌC SINH (XÓA KHOẢNG TRẮNG & VIẾT HOA)",
        goal="Xóa khoảng trắng thừa ở hai đầu chuỗi và viết hoa chữ cái đầu của mỗi từ trong tên.",
        metaphor="Bác thợ cắt tỉa chỉnh sửa lại họ tên cho ngay ngắn trước khi in bằng khen.",
        code_text="""ten_nguoi_dung = "         nguyen van an     "
print("Ten nguoi dung:", ten_nguoi_dung)

ten_da_xoa_khoang_trang = ten_nguoi_dung.strip()
print("Ten nguoi dung sau khi xoa dau cach thua:", ten_da_xoa_khoang_trang)

ten_chuan_hoa = ten_da_xoa_khoang_trang.title()
print("Ten nguoi dung sau khi viet hoa chu cai dau:", ten_chuan_hoa)""",
        output_text="""Ten nguoi dung:          nguyen van an     
Ten nguoi dung sau khi xoa dau cach thua: nguyen van an
Ten nguoi dung sau khi viet hoa chu cai dau: Nguyen Van An""",
        explanations=[
            "ten_nguoi_dung.strip(): Cắt bỏ các khoảng trắng dư thừa ở hai đầu.",
            "ten_da_xoa_khoang_trang.title(): Tự động viết hoa chữ cái đầu mỗi từ."
        ],
        tip_text="Phương thức title() rất thông minh, tự động nhận diện dấu cách để viết hoa chữ tiếp theo."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 1 — BÀI TẬP 3 (GỐC)",
        title="KIỂM TRA ĐỘ DÀI MẬT KHẨU (MẬT KHẨU YẾU / MẠNH)",
        goal="Nhập mật khẩu từ bàn phím, kiểm tra độ dài: < 8 ký tự là yếu, >= 8 ký tự là mạnh.",
        metaphor="Ổ khóa an toàn yêu cầu chìa khóa phải đủ từ 8 răng cưa trở lên mới bảo mật.",
        code_text="""mat_khau = input("Nhap mat khau: ")

if len(mat_khau) < 8:
    print("MAT KHAU QUA YEU")
else:
    print("MAT KHAU MANH")""",
        output_text="""Nhap mat khau: Abc1234
MAT KHAU QUA YEU

Nhap mat khau: Abc12345678
MAT KHAU MANH""",
        explanations=[
            "Hàm len(mat_khau): Trả về tổng số ký tự của mật khẩu vừa nhập.",
            "Cấu trúc if-else phân loại mật khẩu theo ngưỡng 8 ký tự."
        ],
        tip_text="Mật khẩu mạnh thực tế nên có thêm chữ hoa, chữ số và ký tự đặc biệt."
    )

    # 2 bài tập mở rộng Chương 1
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 1 — BÀI TẬP 4 (MỞ RỘNG MỚI)",
        title="ĐẾM SỐ LƯỢNG NGUYÊN ÂM VÀ PHỤ ÂM TRONG CÂU",
        goal="Nhập một câu văn bản tiếng Anh/không dấu, đếm chính xác số nguyên âm (a, e, i, o, u) và số phụ âm.",
        metaphor="Máy quét phân tích ngôn ngữ học bóc tách cấu trúc câu chữ.",
        code_text="""cau = input("Nhap mot cau van: ").lower()

nguyen_am = "aeiou"
so_nguyen_am = 0
so_phu_am = 0

for ch in cau:
    if ch.isalpha():
        if ch in nguyen_am:
            so_nguyen_am += 1
        else:
            so_phu_am += 1

print(f"So luong nguyen am: {so_nguyen_am}")
print(f"So luong phu am: {so_phu_am}")""",
        output_text="""Nhap mot cau van: Lap trinh Python Sao Viet
So luong nguyen am: 7
So luong phu am: 15""",
        explanations=[
            "Chuyển toàn bộ câu sang chữ thường bằng .lower() để so khớp chuẩn xác không phân biệt hoa thường.",
            "Phương thức ch.isalpha() lọc bỏ các ký tự khoảng trắng hoặc dấu câu.",
            "Cú pháp ch in nguyen_am kiểm tra nếu ký tự nằm trong tập 5 nguyên âm 'aeiou'."
        ],
        tip_text="Kỹ thuật đếm tần suất này là nền tảng của các thuật toán xử lý ngôn ngữ tự nhiên (NLP)."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 1 — BÀI TẬP 5 (MỞ RỘNG MỚI)",
        title="ẨN THÔNG TIN SỐ ĐIỆN THOẠI BẢO VỆ DỮ LIỆU CÁ NHÂN",
        goal="Nhập số điện thoại 10 số, che giấu 4 chữ số ở giữa bằng dấu sao '****' (Ví dụ: 0901234567 -> 090****567).",
        metaphor="Tấm rèm bảo mật thông tin tài khoản ngân hàng khi hiển thị trên hóa đơn.",
        code_text="""sdt = input("Nhap so dien thoai (10 so): ").strip()

if len(sdt) == 10 and sdt.isdigit():
    sdt_an = sdt[:3] + "****" + sdt[7:]
    print("So dien thoai sau khi bao mat:", sdt_an)
else:
    print("So dien thoai khong hop le! Vui long nhap dung 10 chu so.")""",
        output_text="""Nhap so dien thoai (10 so): 0901234567
So dien thoai sau khi bao mat: 090****567""",
        explanations=[
            "sdt[:3] lấy 3 chữ số đầu tiên (đầu số nhà mạng).",
            "sdt[7:] lấy 3 chữ số cuối cùng.",
            "Ghép nối chuỗi (String Concatenation) với chuỗi sao '****' ở giữa để bảo vệ quyền riêng tư người dùng."
        ],
        tip_text="Đây là tính năng bảo mật thông tin bắt buộc trong mọi ứng dụng thương mại điện tử hiện đại."
    )


    # ==========================================================================
    # PHẦN 3: CHƯƠNG 2 — KIỂU DỮ LIỆU LIST VÀ DICTIONARY
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 3: CHƯƠNG 2 — KIỂU DỮ LIỆU LIST VÀ DICTIONARY",
        "Học phần Bài 4, Bài 5, 3 bài tập giáo trình gốc và 2 bài tập mở rộng mới."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 4",
        title="KIỂU DỮ LIỆU LIST (DANH SÁCH CÓ THỨ TỰ)",
        goal="Tạo danh sách list, truy cập phần tử qua chỉ số và duyệt danh sách bằng vòng lặp for.",
        metaphor="Ngăn kéo đựng nhiều hộp quà có đánh số thứ tự từ 0, 1, 2, 3...",
        code_text="""a = [2, 5, 10, 1, 3]

print("Phan tu dau tien:", a[0])
print("Phan tu thu hai:", a[1])
print("Phan tu cuoi cung:", a[-1])

print("Duyet danh sach qua tung gia tri:")
for x in a:
    print(x)

print("Duyet danh sach qua vi tri index:")
for i in range(len(a)):
    print(f"Vi tri {i} co gia tri la: {a[i]}")""",
        output_text="""Phan tu dau tien: 2
Phan tu thu hai: 5
Phan tu cuoi cung: 3
Duyet danh sach qua tung gia tri:
2
5
10
1
3
Duyet danh sach qua vi tri index:
Vi tri 0 co gia tri la: 2
Vi tri 1 co gia tri la: 5
Vi tri 2 co gia tri la: 10
Vi tri 3 co gia tri la: 1
Vi tri 4 co gia tri la: 3""",
        explanations=[
            "a = [2, 5, 10, 1, 3]: Khởi tạo một danh sách List chứa 5 số nguyên.",
            "a[0]: Truy cập phần tử đầu tiên tại vị trí 0.",
            "Vòng lặp for x in a: Lần lượt lấy từng giá trị trong danh sách.",
            "Vòng lặp for i in range(len(a)): Duyệt qua từng chỉ số vị trí index."
        ],
        tip_text="List có thể chứa các phần tử có kiểu dữ liệu hỗn hợp (số, chuỗi, boolean)."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 5",
        title="KIỂU DỮ LIỆU DICTIONARY (CẶP KHÓA — GIÁ TRỊ)",
        goal="Nắm vững cú pháp tạo dictionary {key: value}, truy cập qua khóa và duyệt cặp key-value.",
        metaphor="Cuốn danh bạ điện thoại: mỗi cái tên (Key) đi kèm một số điện thoại (Value).",
        code_text="""hoc_sinh = {
    "ten": "Bao Nam",
    "tuoi": 10,
    "lop": "Python Nang Cao"
}

print("Ten hoc sinh:", hoc_sinh["ten"])
print("Tuoi hoc sinh:", hoc_sinh["tuoi"])

print("Duyet qua tung key:")
for k in hoc_sinh:
    print(k, "->", hoc_sinh[k])

print("Duyet qua ca key va value:")
for key, val in hoc_sinh.items():
    print(f"{key}: {val}")""",
        output_text="""Ten hoc sinh: Bao Nam
Tuoi hoc sinh: 10
Duyet qua tung key:
ten -> Bao Nam
tuoi -> 10
lop -> Python Nang Cao
Duyet qua ca key va value:
ten: Bao Nam
tuoi: 10
lop: Python Nang Cao""",
        explanations=[
            "hoc_sinh = {...}: Tạo từ điển với 3 cặp khóa - giá trị.",
            "hoc_sinh['ten']: Truy cập giá trị gắn liền với khóa 'ten'.",
            "hoc_sinh.items(): Trả về toàn bộ các cặp (key, value) để duyệt vòng lặp."
        ],
        tip_text="Khóa (Key) trong Dictionary phải là duy nhất và không được trùng lặp."
    )

    # 3 bài tập gốc Chương 2
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 2 — BÀI TẬP 1 (GỐC)",
        title="QUẢN LÝ THÔNG TIN HỌC SINH BẰNG DICTIONARY",
        goal="Nhập thông tin ten, tuoi, lop từ bàn phím, lưu vào Dictionary và in thông tin định dạng rõ ràng.",
        metaphor="Lập hồ sơ thẻ học sinh điện tử lưu vào hệ thống máy tính của trung tâm.",
        code_text="""ten = input("Nhap ten: ")
tuoi = int(input("Nhap tuoi: "))
lop = input("Nhap lop: ")

hoc_sinh = {
    "ten": ten,
    "tuoi": tuoi,
    "lop": lop
}

print("=== THONG TIN HOC SINH ===")
print("Ten:", hoc_sinh["ten"])
print("Tuoi:", hoc_sinh["tuoi"])
print("Lop:", hoc_sinh["lop"])""",
        output_text="""Nhap ten: Thien
Nhap tuoi: 25
Nhap lop: Python
=== THONG TIN HOC SINH ===
Ten: Thien
Tuoi: 25
Lop: Python""",
        explanations=[
            "Nhận dữ liệu từ bàn phím bằng hàm input() và int(input()).",
            "Đưa các biến vào Dictionary hoc_sinh theo đúng các trường.",
            "Truy xuất và in từng giá trị bằng hoc_sinh['ten'], hoc_sinh['tuoi'], hoc_sinh['lop']."
        ],
        tip_text="Dùng Dictionary giúp dữ liệu có nhãn tên rõ ràng, không bị nhầm lẫn vị trí."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 2 — BÀI TẬP 2 (GỐC)",
        title="TÌM MIN, MAX VÀ SẮP XẾP DANH SÁCH SỐ NGUYÊN",
        goal="Nhập 5 số nguyên vào List, dùng min() và max() tìm giá trị, sắp xếp tăng dần bằng .sort().",
        metaphor="Xếp hàng các bạn nhỏ từ bạn thấp nhất đến bạn cao nhất theo thứ tự đều tăm tắp.",
        code_text="""danh_sach = []

for i in range(1, 6):
    so = int(input(f"Nhap so thu {i}: "))
    danh_sach.append(so)

so_nho_nhat = min(danh_sach)
so_lon_nhat = max(danh_sach)

print(f"So nho nhat: {so_nho_nhat}")
print(f"So lon nhat: {so_lon_nhat}")

danh_sach.sort()
print("Danh sach sap xep tang dan:", danh_sach)""",
        output_text="""Nhap so thu 1: 5
Nhap so thu 2: 1
Nhap so thu 3: 6
Nhap so thu 4: 7
Nhap so thu 5: 10
So nho nhat: 1
So lon nhat: 10
Danh sach sap xep tang dan: [1, 5, 6, 7, 10]""",
        explanations=[
            "danh_sach.append(so): Thêm từng số vừa nhập vào cuối danh sách.",
            "min(danh_sach) và max(danh_sach): Tự động tìm giá trị bé nhất và lớn nhất.",
            "danh_sach.sort(): Sắp xếp trực tiếp các phần tử tăng dần."
        ],
        tip_text="Muốn sắp xếp giảm dần, em truyền tham số danh_sach.sort(reverse=True)."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 2 — BÀI TẬP 3 (GỐC)",
        title="HỆ THỐNG QUẢN LÝ DANH SÁCH MUA ĐỒ CHƠI (LIST & DICT)",
        goal="Xây dựng menu 3 chức năng: 1. Thêm đồ chơi, 2. Xem danh sách, 3. Thoát bằng vòng lặp while True.",
        metaphor="Chiếc máy tính tiền thu ngân trong siêu thị đồ chơi của bé.",
        code_text="""danh_sach_do_choi = []

while True:
    print("=== MENU ===")
    print("1. Them do choi")
    print("2. Xem danh sach")
    print("3. Thoat")
    
    chon = input("Chon chuc nang (1-3): ")
    
    if chon == "1":
        ten = input("Ten mat hang: ")
        so_luong = int(input("So luong: "))
        mon_do = {"ten": ten, "so_luong": so_luong}
        danh_sach_do_choi.append(mon_do)
        print(f"Da them {ten} vao danh sach.\\n")
    elif chon == "2":
        print("--- DANH SACH DO CHOI ---")
        if len(danh_sach_do_choi) == 0:
            print("Danh sach dang trong!\\n")
        else:
            for item in danh_sach_do_choi:
                print(f"- {item['ten']}: {item['so_luong']} cai")
            print()
    elif chon == "3":
        print("Tam biet!")
        break
    else:
        print("Lua chon khong hop le! Vui long chon tu 1 den 3.\\n")""",
        output_text="""=== MENU ===
1. Them do choi
2. Xem danh sach
3. Thoat
Chon chuc nang (1-3): 1
Ten mat hang: robot
So luong: 2
Da them robot vao danh sach.

=== MENU ===
1. Them do choi
2. Xem danh sach
3. Thoat
Chon chuc nang (1-3): 3
Tam biet!""",
        explanations=[
            "Cấu trúc danh sách lồng từ điển: List chứa các Dictionary {'ten': ..., 'so_luong': ...}.",
            "Vòng lặp while True giúp menu tương tác liên tục cho đến khi người dùng chọn phím 3.",
            "Lệnh break kết thúc vòng lặp ngay khi chọn Thoát."
        ],
        tip_text="Đây là mô hình quản lý dữ liệu nền tảng trong mọi phần mềm quản lý kho hàng."
    )

    # 2 bài tập mở rộng Chương 2
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 2 — BÀI TẬP 4 (MỞ RỘNG MỚI)",
        title="TRA CỨU & CẬP NHẬT DANH BẠ ĐIỆN THOẠI BẰNG DICTIONARY",
        goal="Xây dựng từ điển danh bạ gồm tên và số điện thoại, cho phép tra cứu số khi biết tên và cập nhật số mới.",
        metaphor="Ứng dụng danh bạ thông minh trên điện thoại di động.",
        code_text="""danh_ba = {
    "Minh": "0901112233",
    "An": "0904445566",
    "Bao": "0907778899"
}

ten_tim = input("Nhap ten can tra cuu: ")

if ten_tim in danh_ba:
    print(f"So dien thoai cua {ten_tim} la: {danh_ba[ten_tim]}")
else:
    print(f"Khong tim thay {ten_tim} trong danh ba!")
    sdt_moi = input("Nhap so dien thoai de them moi: ")
    danh_ba[ten_tim] = sdt_moi
    print(f"Da them {ten_tim} vao danh ba thanh cong!")""",
        output_text="""Nhap ten can tra cuu: Cuong
Khong tim thay Cuong trong danh ba!
Nhap so dien thoai de them moi: 0912345678
Da them Cuong vao danh ba thanh cong!""",
        explanations=[
            "Toán tử ten_tim in danh_ba kiểm tra sự tồn tại của khóa Key trong từ điển với tốc độ tức thì.",
            "Nếu chưa có, cú pháp danh_ba[ten_tim] = sdt_moi tự động tạo thêm một cặp khóa-giá trị mới."
        ],
        tip_text="Dictionary là cấu trúc dữ liệu tối ưu nhất cho bài toán tìm kiếm và tra cứu theo định danh."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 2 — BÀI TẬP 5 (MỞ RỘNG MỚI)",
        title="THỐNG KÊ TẦN SUẤT XUẤT HIỆN CỦA CÁC TỪ TRONG CÂU",
        goal="Tách câu thành danh sách các từ và dùng Dictionary đếm số lần xuất hiện của từng từ.",
        metaphor="Máy quét kiểm toán từ vựng trong bài văn của học sinh.",
        code_text="""van_ban = "python rat hay va python rat de hoc"
danh_sach_tu = van_ban.split()

thong_ke = {}
for tu in danh_sach_tu:
    if tu in thong_ke:
        thong_ke[tu] += 1
    else:
        thong_ke[tu] = 1

print("=== BANG THONG KE TU VUNG ===")
for tu, so_lan in thong_ke.items():
    print(f"Tu '{tu}': xuat hien {so_lan} lan")""",
        output_text="""=== BANG THONG KE TU VUNG ===
Tu 'python': xuat hien 2 lan
Tu 'rat': xuat hien 2 lan
Tu 'hay': xuat hien 1 lan
Tu 'va': xuat hien 1 lan
Tu 'de': xuat hien 1 lan
Tu 'hoc': xuat hien 1 lan""",
        explanations=[
            "Hàm van_ban.split() tự động tách chuỗi thành List các từ đơn dựa theo dấu khoảng trắng.",
            "Vòng lặp duyệt từng từ: nếu từ đã có trong Dictionary thì tăng số lượng lên 1, nếu chưa có thì gán bằng 1."
        ],
        tip_text="Thuật toán Word Count này là bài toán kinh điển trong lĩnh vực phân tích dữ liệu lớn (Big Data)."
    )


    # ==========================================================================
    # PHẦN 4: CHƯƠNG 3 — ĐỊNH NGHĨA VÀ CÁC THAO TÁC VỚI HÀM
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 4: CHƯƠNG 3 — ĐỊNH NGHĨA VÀ CÁC THAO TÁC VỚI HÀM",
        "Học phần Bài 6, Bài 7, 4 bài tập giáo trình gốc và 2 bài tập mở rộng mới."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 6",
        title="KHÁI NIỆM VÀ VAI TRÒ CỦA HÀM TRONG PYTHON",
        goal="Nắm vững cú pháp tạo hàm bằng từ khóa def và quy tắc thụt lề khối lệnh.",
        metaphor="Robot phụ tá — em dạy việc một lần, khi cần chỉ việc gọi tên Robot!",
        code_text="""def ten_ham(tham_so):
    return gia_tri""",
        output_text="""(Định nghĩa khung mẫu hàm — sẵn sàng gọi trong chương trình)""",
        explanations=[
            "def: Từ khóa bắt buộc để khai báo một hàm mới.",
            "ten_ham: Tên hàm đại diện cho chức năng.",
            "tham_so: Dữ liệu đầu vào hàm nhận để xử lý.",
            "return: Gửi kết quả về cho nơi gọi hàm."
        ],
        tip_text="Nhớ dấu hai chấm (:) ở cuối dòng def và bấm phím Tab lùi đầu dòng khối lệnh."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 7",
        title="ĐỊNH NGHĨA HÀM, GIÁ TRỊ TRẢ VỀ (RETURN) VÀ GỌI HÀM",
        goal="Phân biệt hàm có tham số, hàm trả về kết quả bằng return và hàm trả về None.",
        metaphor="Robot tính 3 x 4 = 12 xong mang số 12 về đưa tận tay cho em cất vào ví.",
        code_text="""def greet(name):
    print("Xin chao!", name)

def multiply(a, b):
    return a * b

greet("Minh")
result = multiply(3, 4)
print("Ket qua phep nhan la:", result)""",
        output_text="""Xin chao! Minh
Ket qua phep nhan la: 12""",
        explanations=[
            "greet('Minh'): Hàm in ra câu chào kèm tên tham số.",
            "multiply(3, 4): Hàm tính tích 2 số và dùng return trả về số 12.",
            "Biến result lưu trữ giá trị trả về để in ra màn hình."
        ],
        tip_text="Khi gặp lệnh return, hàm sẽ kết thúc ngay lập tức."
    )

    # 4 bài tập gốc Chương 3
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 1 (GỐC)",
        title="HÀM KIỂM TRA SỐ CHẴN / SỐ LẺ (IS_EVEN_NUMBER)",
        goal="Viết hàm nhận vào số n, trả về True nếu chẵn, False nếu lẻ.",
        metaphor="Lấy kẹo chia đôi, nếu không dư (n % 2 == 0) là Chẵn, còn dư 1 là Lẻ.",
        code_text="""def is_even_number(n):
    if n % 2 == 0:
        return True
    else:
        return False

n = int(input("Nhap so nguyen duong n: "))
if is_even_number(n):
    print(n, "la so chan")
else:
    print(n, "la so le")""",
        output_text="""Nhap so nguyen duong n: 4
4 la so chan""",
        explanations=[
            "Toán tử n % 2 là phép chia lấy phần dư.",
            "Hàm trả về giá trị Boolean (True / False) để câu lệnh if-else sử dụng."
        ],
        tip_text="Có thể viết gọn hàm chỉ bằng một dòng: return n % 2 == 0."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 2 (GỐC)",
        title="HÀM TÍNH TỔNG TỪ 1 ĐẾN N (SUM_FROM_1_TO_N)",
        goal="Viết hàm tính tổng: 1 + 2 + 3 + ... + n. Nếu n <= 0 trả về 0.",
        metaphor="Bỏ từng đồng xu từ 1 đến n vào heo đất tiết kiệm rồi đếm tổng số xu.",
        code_text="""def sum_from_1_to_n(n):
    if n <= 0:
        return 0
    total = 0
    for i in range(1, n + 1):
        total = total + i
    return total

n = int(input("Nhap so nguyen duong n: "))
result = sum_from_1_to_n(n)
print(f"Tong tu 1 den {n} la: {result}")""",
        output_text="""Nhap so nguyen duong n: 10
Tong tu 1 den 10 la: 55""",
        explanations=[
            "Nếu n <= 0, hàm lập tức trả về 0.",
            "Vòng lặp for i in range(1, n + 1) cộng dồn từng số vào biến total."
        ],
        tip_text="range(1, n + 1) cần có n + 1 để lấy được chính xác số n."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 3 (GỐC)",
        title="HÀM IN BẢNG CỬU CHƯƠNG (PRINT_MULTIPLICATION_TABLE)",
        goal="Viết hàm in ra bảng cửu chương của số n từ 1 đến 10.",
        metaphor="Chiếc máy in thông minh tự động in 10 dòng phép nhân.",
        code_text="""def print_multiplication_table(n):
    print(f"Bang cuu chuong cua {n} la:")
    for i in range(1, 11):
        print(f"{n} x {i} = {n * i}")

n = int(input("Nhap so nguyen duong n: "))
print_multiplication_table(n)""",
        output_text="""Nhap so nguyen duong n: 5
Bang cuu chuong cua 5 la:
5 x 1 = 5
5 x 2 = 10
...
5 x 9 = 45
5 x 10 = 50""",
        explanations=[
            "Vòng lặp for i in range(1, 11) chạy qua 10 số từ 1 đến 10.",
            "Dùng f-string f'{n} x {i} = {n * i}' để in và tính phép nhân."
        ],
        tip_text="Hàm này chỉ làm nhiệm vụ in nên không cần câu lệnh return."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 4 (GỐC)",
        title="DỰ ÁN MÁY TÍNH MINI 4 PHÉP TÍNH (CỘNG - TRỪ - NHÂN - CHIA)",
        goal="Viết 4 hàm toán học và menu 1-4, kiểm tra điều kiện chia cho 0 và lặp lại khi nhập sai.",
        metaphor="Tự chế tạo chiếc máy tính bỏ túi Casio tí hon bằng các hàm riêng biệt.",
        code_text="""def sum(a, b):
    return a + b

def subtract(a, b):
    return a - b

def multiply(a, b):
    return a * b

def divide(a, b):
    if b == 0:
        return "Khong the chia cho 0!"
    return a / b

while True:
    print("1 - Cong")
    print("2 - Tru")
    print("3 - Nhan")
    print("4 - Chia")
    choice = int(input("Hay chon phep tinh (1-4): "))
    if 1 <= choice <= 4:
        a = float(input("Nhap so a: "))
        b = float(input("Nhap so b: "))
        if choice == 1:
            print("Ket qua:", sum(a, b))
        elif choice == 2:
            print("Ket qua:", subtract(a, b))
        elif choice == 3:
            print("Ket qua:", multiply(a, b))
        elif choice == 4:
            print("Ket qua:", divide(a, b))
        break
    else:
        print("Khong ton tai phep tinh! Vui long nhap lai.\\n")""",
        output_text="""1 - Cong
2 - Tru
3 - Nhan
4 - Chia
Hay chon phep tinh (1-4): 2
Nhap so a: 10
Nhap so b: 15
Ket qua: -5.0""",
        explanations=[
            "Chương trình chia thành 4 hàm độc lập (sum, subtract, multiply, divide).",
            "Kiểm tra điều kiện b == 0 trong hàm divide để chống lỗi sập chương trình.",
            "Vòng lặp while True kết hợp break giúp xử lý nhập sai linh hoạt."
        ],
        tip_text="Dùng float(input()) để tính toán được cả các số thập phân."
    )

    # 2 bài tập mở rộng Chương 3
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 5 (MỞ RỘNG MỚI)",
        title="HÀM KIỂM TRA SỐ ĐỐI XỨNG (PALINDROMIC NUMBER)",
        goal="Viết hàm kiem_tra_doi_xung(n) trả về True nếu đọc xuôi đọc ngược đều bằng nhau (ví dụ: 121, 1331).",
        metaphor="Tấm gương soi phản chiếu hai chiều giống hệt nhau.",
        code_text="""def kiem_tra_doi_xung(n):
    chuoi_so = str(n)
    return chuoi_so == chuoi_so[::-1]

n = int(input("Nhap so nguyen duong can kiem tra: "))
if kiem_tra_doi_xung(n):
    print(f"So {n} la so doi xung!")
else:
    print(f"So {n} khong phai la so doi xung.")""",
        output_text="""Nhap so nguyen duong can kiem tra: 12321
So 12321 la so doi xung!""",
        explanations=[
            "Chuyển số n thành chuỗi str(n) để áp dụng kỹ thuật đảo ngược slicing [::-1].",
            "So sánh chuỗi gốc với chuỗi đảo: nếu bằng nhau thì trả về True, ngược lại False."
        ],
        tip_text="Kỹ thuật này áp dụng được cho cả chuỗi từ đối xứng (ví dụ: 'radar', 'level')."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 3 — BÀI TẬP 6 (MỞ RỘNG MỚI)",
        title="HÀM TÍNH CHỈ SỐ KHỐI CƠ THỂ BMI VÀ TƯ VẤN SỨC KHỎE",
        goal="Viết hàm tinh_bmi(can_nang, chieu_cao) và in kết quả phân loại: Gầy (< 18.5), Bình thường (18.5 - 24.9), Thừa cân (>= 25).",
        metaphor="Bác sĩ điện tử thông minh đo khám sức khỏe học đường cho bé.",
        code_text="""def tinh_bmi(can_nang, chieu_cao):
    bmi = can_nang / (chieu_cao ** 2)
    return bmi

can_nang = float(input("Nhap can nang (kg): "))
chieu_cao = float(input("Nhap chieu cao (m, vi du 1.45): "))

chi_so = tinh_bmi(can_nang, chieu_cao)
print(f"Chi so BMI cua be la: {chi_so:.1f}")

if chi_so < 18.5:
    print("Danh gia: The trang hoi gay, can bo sung dinh duong!")
elif 18.5 <= chi_so < 25:
    print("Danh gia: The trang ly tuong rat can doi!")
else:
    print("Danh gia: The trang hoi thua can, hay cham tap the thao!")""",
        output_text="""Nhap can nang (kg): 35
Nhap chieu cao (m, vi du 1.45): 1.40
Chi so BMI cua be la: 17.9
Danh gia: The trang hoi gay, can bo sung dinh duong!""",
        explanations=[
            "Công thức chuẩn: BMI = Cân nặng (kg) / (Chiều cao (m))².",
            "Hàm trả về kết quả số thực và dùng f-string {chi_so:.1f} làm tròn 1 chữ số thập phân."
        ],
        tip_text="Hàm giúp chia nhỏ phép tính toán học phức tạp thành module độc lập, dễ kiểm tra."
    )


    # ==========================================================================
    # PHẦN 5: CHƯƠNG 4 — SỬ DỤNG CÁC THƯ VIỆN CƠ BẢN CÓ SẴN
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 5: CHƯƠNG 4 — SỬ DỤNG CÁC THƯ VIỆN CƠ BẢN (RANDOM, MATH, DATETIME)",
        "Học phần Bài 8, Bài 9, Bài 10, Bài 11, 4 bài tập giáo trình gốc và 2 bài tập mở rộng mới."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 8",
        title="GIỚI THIỆU THƯ VIỆN VÀ CÁC CÁCH IMPORT TRONG PYTHON",
        goal="Hiểu module là gì và 3 cách import: import toàn bộ, from ... import và đặt tên rút gọn as.",
        metaphor="Hộp đồ nghề Doraemon — có sẵn bảo bối, em chỉ cần lấy ra dùng.",
        code_text="""import math
print("Can bac hai cua 16 la:", math.sqrt(16))

from math import sqrt
print("Can bac hai cua 25 la:", sqrt(25))

import datetime as dt
now = dt.datetime.now()
print("Thoi gian hien tai:", now)""",
        output_text="""Can bac hai cua 16 la: 4.0
Can bac hai cua 25 la: 5.0
Thoi gian hien tai: 2026-08-28 20:00:00.123456""",
        explanations=[
            "import math: Nhập toàn bộ thư viện, gọi qua math.sqrt().",
            "from math import sqrt: Chỉ lấy hàm sqrt để gọi trực tiếp.",
            "import datetime as dt: Đặt tên tắt dt cho ngắn gọn."
        ],
        tip_text="Nên đặt tất cả các lệnh import ở ngay các dòng đầu tiên của file."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 9",
        title="THƯ VIỆN RANDOM — LÀM VIỆC VỚI SỐ NGẪU NHIÊN",
        goal="Dùng random.randint(a, b) sinh số nguyên và random.choice(list) chọn phần tử ngẫu nhiên.",
        metaphor="Chiếc nón kỳ diệu bốc thăm trúng thưởng, mỗi lần rút ra một món quà bất ngờ.",
        code_text="""import random

so_ngau_nhien = random.randint(1, 10)
print("So may man hom nay la:", so_ngau_nhien)

hoa_qua = ["Tao", "Chuoi", "Cam", "Xoai"]
mon_duoc_chon = random.choice(hoa_qua)
print("Hom nay be duoc an qua:", mon_duoc_chon)""",
        output_text="""So may man hom nay la: 7
Hom nay be duoc an qua: Cam""",
        explanations=[
            "random.randint(1, 10): Sinh số nguyên ngẫu nhiên từ 1 đến 10.",
            "random.choice(hoa_qua): Chọn ngẫu nhiên một phần tử trong danh sách."
        ],
        tip_text="Mỗi lần chạy kết quả sẽ khác nhau vì do máy tính tự quay số."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 10",
        title="THƯ VIỆN MATH — CÁC PHÉP TOÁN NÂNG CAO",
        goal="Sử dụng math.pi, math.sqrt(), math.pow() và math.factorial().",
        metaphor="Compa và máy tính khoa học giúp bé giải hình học cực nhanh.",
        code_text="""import math

print("So Pi =", math.pi)
print("Can bac hai cua 49 la:", math.sqrt(49))
print("2 mu 5 la:", math.pow(2, 5))
print("5 giai thua (5!) la:", math.factorial(5))""",
        output_text="""So Pi = 3.141592653589793
Can bac hai cua 49 la: 7.0
2 mu 5 la: 32.0
5 giai thua (5!) la: 120""",
        explanations=[
            "math.pi: Hằng số số Pi xấp xỉ 3.14159...",
            "math.sqrt(49): Tính căn bậc hai của 49 (ra 7.0).",
            "math.pow(2, 5): Tính lũy thừa 2 mũ 5 (ra 32.0).",
            "math.factorial(5): Tính giai thừa 5! = 1*2*3*4*5 = 120."
        ],
        tip_text="Kết quả của sqrt() và pow() luôn luôn là kiểu số thực float."
    )

    add_lesson_box_bw(
        doc,
        badge="BÀI 11",
        title="THƯ VIỆN DATETIME — LÀM VIỆC VỚI NGÀY VÀ GIỜ",
        goal="Lấy thời gian thực của máy tính và bóc tách Ngày, Tháng, Năm, Giờ, Phút.",
        metaphor="Chiếc đồng hồ thông minh gắn trên máy tính báo cho chương trình biết giờ giấc.",
        code_text="""import datetime

now = datetime.datetime.now()

print("Bay gio la:", now.hour, "gio", now.minute, "phut")
print("Hom nay la ngay:", now.day, "thang", now.month, "nam", now.year)""",
        output_text="""Bay gio la: 20 gio 0 phut
Hom nay la ngay: 28 thang 8 nam 2026""",
        explanations=[
            "datetime.datetime.now(): Lấy thời gian hiện tại của hệ thống.",
            "now.hour, now.minute: Lấy giờ và phút.",
            "now.day, now.month, now.year: Lấy ngày, tháng và năm."
        ],
        tip_text="Nếu chỉ muốn lấy ngày mà không cần giờ, dùng datetime.date.today()."
    )

    # 4 bài tập gốc Chương 4
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 1 (GỐC)",
        title="MÔ PHỎNG TUNG XÚC XẮC N LẦN (DÙNG RANDOM)",
        goal="Mô phỏng tung xúc xắc 6 mặt n lần và in ra tổng số điểm sau n lần tung.",
        metaphor="Bé lắc xí ngầu từ 1 đến 6 chấm trong trò chơi cờ cá ngựa.",
        code_text="""import random

n = int(input("Nhap so lan tung: "))
tong_diem = 0

for i in range(1, n + 1):
    diem = random.randint(1, 6)
    print(f"Lan tung thu {i} duoc diem: {diem}")
    tong_diem = tong_diem + diem

print(f"Tong diem sau {n} lan tung: {tong_diem}")""",
        output_text="""Nhap so lan tung: 3
Lan tung thu 1 duoc diem: 5
Lan tung thu 2 duoc diem: 6
Lan tung thu 3 duoc diem: 4
Tong diem sau 3 lan tung: 15""",
        explanations=[
            "Mỗi lần lặp, random.randint(1, 6) sinh điểm ngẫu nhiên.",
            "Cộng dồn điểm vào biến tong_diem và in kết quả cuối cùng."
        ],
        tip_text="Đề bài yêu cầu chỉ in kết quả ra màn hình mà không cần return."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 2 (GỐC)",
        title="TÍNH DIỆN TÍCH, CHU VI & THỂ TÍCH HÌNH HỌC (DÙNG MATH)",
        goal="Viết 3 hàm riêng biệt cho hình tròn (r), hình cầu (r) và hình trụ (r, h) theo đúng công thức.",
        metaphor="Kiến trúc sư tính toán kích thước quả bóng tròn và lon nước ngọt hình trụ.",
        code_text="""import math

def hinh_tron(r):
    chu_vi = 2 * math.pi * r
    dien_tich = math.pi * (r ** 2)
    print(f"Hinh tron ban kinh r = {r} , chu vi = {chu_vi:.2f} , dien tich = {dien_tich:.2f}")

def hinh_cau(r):
    the_tich = (4/3) * math.pi * (r ** 3)
    print(f"Hinh cau ban kinh r = {r} , the tich = {the_tich:.2f}")

def hinh_tru(r, h):
    the_tich = math.pi * (r ** 2) * h
    print(f"Hinh tru ban kinh r = {r} , chieu cao h = {h} , the tich = {the_tich:.2f}")

hinh_tron(4)
hinh_cau(4)
hinh_tru(4, 10)""",
        output_text="""Hinh tron ban kinh r = 4 , chu vi = 25.13 , dien tich = 50.27
Hinh cau ban kinh r = 4 , the tich = 268.08
Hinh tru ban kinh r = 4 , chieu cao h = 10 , the tich = 502.65""",
        explanations=[
            "Hàm hinh_tron(r): Chu vi = 2πr, Diện tích = πr².",
            "Hàm hinh_cau(r): Thể tích cầu = (4/3)πr³.",
            "Hàm hinh_tru(r, h): Thể tích trụ = πr²h.",
            "Dùng định dạng :.2f để làm tròn đúng 2 chữ số thập phân."
        ],
        tip_text="Trong Python, phép tính lũy thừa viết là r ** 2 hoặc r ** 3."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 3 (GỐC)",
        title="HIỂN THỊ THỜI GIAN THEO ĐỊNH DẠNG DỄ ĐỌC (DATETIME)",
        goal="Lấy thời gian hiện tại bằng datetime.datetime.now() và in định dạng theo mẫu giáo trình.",
        metaphor="Bản tin dự báo thời tiết đọc thông báo giờ giấc chính xác cho mọi người.",
        code_text="""import datetime

now = datetime.datetime.now()
print(f"Bay gio la {now.hour} : {now.minute} ngay {now.day} / {now.month} / {now.year}")""",
        output_text="""Bay gio la 9 : 36 ngay 22 / 12 / 2025""",
        explanations=[
            "Lấy thời gian thực của máy tính bằng datetime.datetime.now().",
            "Dùng f-string sắp xếp thứ tự giờ : phút và ngày / tháng / năm."
        ],
        tip_text="Có thể mở rộng thêm giây bằng cú pháp {now.second}."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 4 (GỐC)",
        title="TRÒ CHƠI ĐOÁN SỐ BÍ MẬT TỪ 1 ĐẾN 100 (RANDOM MINI GAME)",
        goal="Tạo số bí mật bằng random.randint(1, 100), nhận số đoán và gợi ý Quá cao / Quá thấp.",
        metaphor="Trò chơi 'Đi tìm kho báu' — máy tính là người quản trò thông thái dẫn đường cho bé.",
        code_text="""import random

so_bi_mat = random.randint(1, 100)

print("====================================")
print("Toi da nghi ra mot so tu 1 den 100.")
print("Ban hay doan xem do la so gi!")
print("====================================")

while True:
    doan = int(input("Nhap so cua ban: "))
    
    if doan == so_bi_mat:
        print("Chuc mung! Ban da doan dung.")
        break
    elif doan < so_bi_mat:
        print("So cua ban thap hon so bi mat. Hay thu so cao hon!")
    else:
        print("So cua ban cao hon so bi mat. Hay thu so thap hon!")""",
        output_text="""====================================
Toi da nghi ra mot so tu 1 den 100.
Ban hay doan xem do la so gi!
====================================
Nhap so cua ban: 50
So cua ban thap hon so bi mat. Hay thu so cao hon!
Nhap so cua ban: 75
So cua ban cao hon so bi mat. Hay thu so thap hon!
Nhap so cua ban: 51
Chuc mung! Ban da doan dung.""",
        explanations=[
            "random.randint(1, 100): Chọn ngẫu nhiên số bí mật từ 1 đến 100.",
            "Vòng lặp while True nhận số đoán cho đến khi đoán đúng.",
            "Lệnh break kích hoạt khi doan == so_bi_mat để kết thúc game."
        ],
        tip_text="Mẹo chiến thắng nhanh: Luôn đoán số ở chính giữa khoảng (Ví dụ 50 trước) để thu hẹp phạm vi nhanh nhất!"
    )

    # 2 bài tập mở rộng Chương 4
    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 5 (MỞ RỘNG MỚI)",
        title="MÁY TẠO MẬT KHẨU BẢO MẬT NGẪU NHIÊN (PASSWORD GENERATOR)",
        goal="Tự động sinh chuỗi mật khẩu ngẫu nhiên có độ dài n gồm chữ hoa, chữ thường và chữ số bằng thư viện random.",
        metaphor="Chiếc máy cấp mã khóa OTP ngân hàng tự động bảo mật tuyệt đối.",
        code_text="""import random

ky_tu = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789!@#$%^&*"
do_dai = int(input("Do dai mat khau mong muon (vi du 8, 12): "))

mat_khau_moi = ""
for i in range(do_dai):
    mat_khau_moi += random.choice(ky_tu)

print("Mat khau ngau nhien an toan cua ban la:", mat_khau_moi)""",
        output_text="""Do dai mat khau mong muon (vi du 8, 12): 10
Mat khau ngau nhien an toan cua ban la: K9#mP2$xR7""",
        explanations=[
            "random.choice(ky_tu) chọn ngẫu nhiên một ký tự từ kho chuỗi ký tự an toàn.",
            "Vòng lặp chạy n lần ghép nối các ký tự lại với nhau thành mật khẩu hoàn chỉnh."
        ],
        tip_text="Mật khẩu ngẫu nhiên này có độ bảo mật cực cao, chống lại các phần mềm dò quét mật khẩu."
    )

    add_lesson_box_bw(
        doc,
        badge="CHƯƠNG 4 — BÀI TẬP 6 (MỞ RỘNG MỚI)",
        title="ĐỒNG HỒ ĐẾM NGƯỢC THỜI GIAN ĐẾN NGÀY TẾT NGUYÊN ĐÁN",
        goal="Sử dụng module datetime tính toán chính xác còn bao nhiêu ngày, giờ, phút nữa là đến thời khắc năm mới.",
        metaphor="Đồng hồ đếm ngược giao thừa tại Quảng trường thành phố.",
        code_text="""import datetime

hom_nay = datetime.datetime.now()
ngay_tet = datetime.datetime(2027, 1, 1, 0, 0, 0)

khoang_cach = ngay_tet - hom_nay

print(f"Thoi diem hien tai: {hom_nay.strftime('%d/%m/%Y %H:%M:%S')}")
print(f"Con dung: {khoang_cach.days} ngay nua la den Tet!")""",
        output_text="""Thoi diem hien tai: 28/08/2026 20:15:00
Con dung: 125 ngay nua la den Tet!""",
        explanations=[
            "Phép trừ giữa 2 đối tượng datetime tạo ra một đối tượng khoảng thời gian timedelta.",
            "khoang_cach.days trả về chính xác số ngày chênh lệch giữa 2 mốc thời gian."
        ],
        tip_text="strftime('%d/%m/%Y %H:%M:%S') là phương thức định dạng ngày tháng năm giờ phút giây chuẩn quốc tế."
    )


    # ==========================================================================
    # PHẦN 6: BÀI TẬP PYTHON NÂNG CAO TỔNG HỢP (6 DỰ ÁN LIÊN MÔN)
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 6: BÀI TẬP PYTHON NÂNG CAO TỔNG HỢP (6 DỰ ÁN NÂNG CAO)",
        "Bộ bài tập dự án tích hợp Hàm (def), Danh sách (List), Từ điển (Dictionary), Xử lý chuỗi và Đồ họa Turtle."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 1",
        title="HÀM VẼ CHUỖI HÌNH VUÔNG NẰM NGANG TĂNG DẦN KÍCH THƯỚC (VE_CHUOI_HINH_VUONG)",
        goal="Xây dựng hàm nhận vào List kích thước cạnh và khoảng cách dịch chuyển, vẽ các hình vuông nối tiếp nhau.",
        metaphor="Xây dựng dãy nhà phố liên kế có diện tích mở rộng dần theo chiều ngang quy hoạch.",
        code_text="""import turtle

def ve_chuoi_hinh_vuong(danh_sach_canh, khoang_cach):
    print("Dang khoi chay giao dien Turtle de ve chuoi hinh vuong...")
    but_ve = turtle.Turtle()
    but_ve.pensize(2)
    but_ve.speed(3)
    
    for canh in danh_sach_canh:
        for i in range(4):
            but_ve.forward(canh)
            but_ve.right(90)
        but_ve.penup()
        but_ve.forward(canh + khoang_cach)
        but_ve.pendown()
    
    turtle.done()

n = int(input("So luong hinh vuong muon ve: "))
danh_sach_canh = []
for i in range(1, n + 1):
    canh = int(input(f"Kich thuoc canh hinh thu {i}: "))
    danh_sach_canh.append(canh)

khoang_cach = int(input("Khoang cach dich chuyen giua cac hinh: "))
ve_chuoi_hinh_vuong(danh_sach_canh, khoang_cach)""",
        output_text="""So luong hinh vuong muon ve: 3
Kich thuoc canh hinh thu 1: 40
Kich thuoc canh hinh thu 2: 60
Kich thuoc canh hinh thu 3: 80
Khoang cach dich chuyen giua cac hinh: 20
Dang khoi chay giao dien Turtle de ve chuoi hinh vuong...
(Cửa sổ Turtle hiển thị 3 hình vuông nằm ngang lần lượt có cạnh 40, 60, 80 cách nhau 20 bước)""",
        explanations=[
            "Hàm ve_chuoi_hinh_vuong nhận danh sách cạnh và khoảng cách dịch chuyển.",
            "Vòng lặp ngoài duyệt từng kích thước; vòng lặp trong 4 lần vẽ hình vuông.",
            "Sau mỗi hình, nhấc bút penup(), tiến tới đoạn (canh + khoang_cach), rồi hạ bút pendown()."
        ],
        tip_text="Dùng append() để thu thập các kích thước cạnh người dùng nhập vào mảng List trước khi gọi hàm."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 2",
        title="HÀM VẼ ĐA GIÁC HOA VĂN MÀU SẮC ĐỐI XỨNG (VE_DA_GIAC_MAU_SAC)",
        goal="Duyệt Dictionary chứa cặp màu sắc và bán kính, dùng .strip() làm sạch chuỗi màu và vẽ bông hoa nhiều cánh xoay quanh tâm.",
        metaphor="Bông hoa đa giác nhiều tầng nở rộ với các cánh hoa tròn đều tăm tắp bao quanh nhụy hoa.",
        code_text="""import turtle

def ve_da_giac_mau_sac(tu_dien_mau, so_hinh_tron):
    print("Dang mo giao dien Turtle de ve da giac mau sac...")
    but_ve = turtle.Turtle()
    but_ve.pensize(2)
    but_ve.speed(4)
    goc_quay = 360 / so_hinh_tron
    
    for mau, ban_kinh in tu_dien_mau.items():
        mau_sach = mau.strip()
        but_ve.color(mau_sach)
        for i in range(so_hinh_tron):
            but_ve.circle(ban_kinh)
            but_ve.left(goc_quay)
    
    turtle.done()

tu_dien_mau = {" red ": 50, " blue ": 80, " green ": 110}
so_hinh_tron = int(input("Nhap so luong hinh tron muon lap lai tai moi tam: "))
ve_da_giac_mau_sac(tu_dien_mau, so_hinh_tron)""",
        output_text="""Nhap so luong hinh tron muon lap lai tai moi tam: 4
Dang mo giao dien Turtle de ve da giac mau sac...
(Cửa sổ Turtle vẽ 3 cụm hoa văn 4 cánh lồng nhau tại tâm với 3 màu Đỏ - Xanh dương - Xanh lá, bán kính 50, 80, 110)""",
        explanations=[
            "tu_dien_mau.items(): Lấy ra từng cặp (màu sắc, bán kính) từ Dictionary.",
            "mau.strip(): Cắt bỏ các khoảng trắng thừa ở hai đầu chuỗi màu tránh lỗi cho lệnh .color().",
            "goc_quay = 360 / so_hinh_tron: Chia đều góc xoay để cánh hoa phân bố đối xứng hoàn hảo."
        ],
        tip_text="Phương thức .strip() là kỹ thuật cơ bản giúp làm sạch dữ liệu đầu vào cực kỳ hữu ích."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 3",
        title="HÀM QUẢN LÝ VÀ XẾP LOẠI ĐIỂM HỌC SINH (LAP_DANH_SACH_LOP)",
        goal="Nhập điểm cho danh sách học sinh có sẵn, lưu vào Dictionary và phân loại học lực Giỏi (>= 8.0) hoặc Đạt (< 8.0).",
        metaphor="Hệ thống sổ điểm điện tử của giáo viên, tự động tổng hợp và xếp loại thi đua học kỳ.",
        code_text="""def lap_danh_sach_lop(danh_sach_ten):
    bang_diem = {}
    for ten in danh_sach_ten:
        diem = float(input(f"Nhap diem cho {ten}: "))
        bang_diem[ten] = diem
    
    print("\\n=== KET QUA XEP LOAI HOC LUC ===")
    for ten, diem in bang_diem.items():
        if diem >= 8.0:
            xep_loai = "Gioi"
        else:
            xep_loai = "Dat"
        print(f"Hoc sinh {ten} - Diem: {diem} - Xep loai: {xep_loai}")

danh_sach_hoc_sinh = ["Nguyen Tri Dung", "Tran Thu Ha", "Le Minh Khoi"]
lap_danh_sach_lop(danh_sach_hoc_sinh)""",
        output_text="""Nhap diem cho Nguyen Tri Dung: 8.5
Nhap diem cho Tran Thu Ha: 7.2
Nhap diem cho Le Minh Khoi: 9.0

=== KET QUA XEP LOAI HOC LUC ===
Hoc sinh Nguyen Tri Dung - Diem: 8.5 - Xep loai: Gioi
Hoc sinh Tran Thu Ha - Diem: 7.2 - Xep loai: Dat
Hoc sinh Le Minh Khoi - Diem: 9.0 - Xep loai: Gioi""",
        explanations=[
            "Hàm lap_danh_sach_lop duyệt danh sách họ tên và gọi hàm input() nhập điểm.",
            "float(input()): Ép kiểu dữ liệu điểm số sang dạng số thực.",
            "bang_diem[ten] = diem: Lưu thông tin vào Dictionary với tên làm Key và điểm làm Value.",
            "Cấu trúc if-else so sánh điểm >= 8.0 để gán xếp loại."
        ],
        tip_text="Dùng Dictionary giúp tra cứu điểm của từng học sinh theo tên tức thì."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 4",
        title="HÀM KẾT HỢP LIST VẼ BIỂU ĐỒ CỘT TĂNG TRƯỞNG TURTLE (VE_BIEU_DO_TANG_TRUONG)",
        goal="Nhận danh sách chiều cao cột và màu sắc, vẽ biểu đồ cột hình chữ nhật có khoảng cách trống 15 bước giữa các cột.",
        metaphor="Biểu đồ thống kê kết quả học tập và phát triển của học viên trên màn hình trực quan.",
        code_text="""import turtle

def ve_bieu_do_tang_truong(danh_sach_chieu_cao, mau_sac):
    print("Dang ve bieu do cot tang truong tren Turtle...")
    but_ve = turtle.Turtle()
    but_ve.pensize(2)
    but_ve.color(mau_sac)
    but_ve.speed(3)
    
    do_rong = 30
    for h in danh_sach_chieu_cao:
        but_ve.forward(do_rong)
        but_ve.left(90)
        but_ve.forward(h)
        but_ve.left(90)
        but_ve.forward(do_rong)
        but_ve.left(90)
        but_ve.forward(h)
        but_ve.left(90)
        
        but_ve.penup()
        but_ve.forward(do_rong + 15)
        but_ve.pendown()
        
    turtle.done()

danh_sach_chieu_cao = [50, 120, 80, 150, 95]
mau = input("Nhap mau sac cho bieu do tu ban phim: ")
ve_bieu_do_tang_truong(danh_sach_chieu_cao, mau)""",
        output_text="""Nhap mau sac cho bieu do tu ban phim: orange
Dang ve bieu do cot tang truong tren Turtle...
(Cửa sổ Turtle vẽ trục biểu đồ gồm 5 cột đứng màu cam có độ cao 50, 120, 80, 150, 95 cách đều nhau 15 bước)""",
        explanations=[
            "Hàm ve_bieu_do_tang_truong duyệt qua từng chiều cao h trong danh sách.",
            "Thuật toán vẽ 1 cột đứng: đi ngang do_rong (30 bước), đi lên h bước, đi ngang do_rong, đi xuống h bước.",
            "Nhấc bút penup(), tiến tới đoạn (do_rong + 15 bước) để tạo khoảng cách đều giữa các cột."
        ],
        tip_text="Đây là ứng dụng trực quan hóa dữ liệu (Data Visualization) thực tế bằng đồ họa Python."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 5",
        title="HÀM BỘ LỌC AN NINH QUÉT MÃ ĐĂNG KÝ VÀ CẬP NHẬT THÀNH VIÊN CLB (BO_LOC_DANG_KY_CLB)",
        goal="Quét danh sách mã, làm sạch bằng .strip(), kiểm tra .isalnum() loại bỏ mã lỗi và nhập tên thành viên cho mã hợp lệ.",
        metaphor="Cổng an ninh soát vé thông minh tự động quét mã thẻ thành viên câu lạc bộ trường học.",
        code_text="""def bo_loc_dang_ky_clb(danh_sach_ma):
    thanh_vien = {}
    print("=== TIEN TRINH QUET MA DANG KY CLB ===")
    for ma in danh_sach_ma:
        ma_sach = ma.strip()
        if ma_sach.isalnum():
            ten = input(f"Ma hop le [{ma_sach}] -> Nhap ho ten hoc sinh: ")
            thanh_vien[ma_sach] = ten
        else:
            print(f"Ma khong hop le [{ma_sach}] -> Da bi loai tu dong!")
            
    print("\\n=== DANH SACH THANH VIEN CLB CHINH THUC ===")
    for ma, ten in thanh_vien.items():
        print(f"- Ma: {ma} | Ho ten: {ten}")

danh_sach_ma = ["THPT2026", "KHTN 123", "CHUYEN_ANH", "TIN2K9"]
bo_loc_dang_ky_clb(danh_sach_ma)""",
        output_text="""=== TIEN TRINH QUET MA DANG KY CLB ===
Ma hop le [THPT2026] -> Nhap ho ten hoc sinh: Nguyen Hoang Long
Ma khong hop le [KHTN 123] -> Da bi loai tu dong!
Ma khong hop le [CHUYEN_ANH] -> Da bi loai tu dong!
Ma hop le [TIN2K9] -> Nhap ho ten hoc sinh: Phan My Linh

=== DANH SACH THANH VIEN CLB CHINH THUC ===
- Ma: THPT2026 | Ho ten: Nguyen Hoang Long
- Ma: TIN2K9 | Ho ten: Phan My Linh""",
        explanations=[
            "ma.strip(): Cắt bỏ các ký tự khoảng trắng thừa ở hai đầu chuỗi mã.",
            "ma_sach.isalnum(): Kiểm tra chuỗi chỉ gồm chữ cái và số. Các mã có khoảng trắng ('KHTN 123') hoặc gạch dưới ('CHUYEN_ANH') bị loại tự động.",
            "Đối với mã hợp lệ, kích hoạt input() nhập tên và lưu cặp {ma: ten} vào Dictionary thành viên."
        ],
        tip_text="Phương thức isalnum() viết tắt của 'is alpha-numeric' (kiểm tra chữ và số kết hợp)."
    )

    add_lesson_box_bw(
        doc,
        badge="DỰ ÁN TỔNG HỢP 6",
        title="HÀM VẼ HOA VĂN XOAY VÒNG TỪ ĐA GIÁC ĐỀU (VE_HOA_VAN_XOAY)",
        goal="Xây dựng hàm nhận vào số cạnh, chiều dài cạnh, độ dày nét vẽ; dùng vòng lặp lồng nhau xoay 6 lần tạo hoa văn 6 cánh.",
        metaphor="Kính vạn hoa nghệ thuật tạo nên những đóa hoa đối xứng hình học kỳ diệu.",
        code_text="""import turtle

def ve_hoa_van_xoay(so_canh, chieu_dai_canh, do_day):
    print("Dang khoi tao thuat toan xoay da giac de tao hoa van nghe thuat...")
    but_ve = turtle.Turtle()
    but_ve.pensize(do_day)
    but_ve.color("purple")
    but_ve.speed(5)
    
    goc_da_giac = 360 / so_canh
    goc_xoay_truc = 360 / 6
    
    for i in range(6):
        for j in range(so_canh):
            but_ve.forward(chieu_dai_canh)
            but_ve.right(goc_da_giac)
        but_ve.right(goc_xoay_truc)
        
    turtle.done()

so_canh = int(input("Em muon canh hoa la hinh gi? (Nhap 3: Tam giac, 4: Hinh vuong, 5: Ngu giac): "))
chieu_dai = int(input("Nhap chieu dai mot canh cua da giac: "))
do_day = int(input("Nhap do day net ve (tu 1 den 5): "))

ve_hoa_van_xoay(so_canh, chieu_dai, do_day)""",
        output_text="""Em muon canh hoa la hinh gi? (Nhap 3: Tam giac, 4: Hinh vuong, 5: Ngu giac): 4
Nhap chieu dai mot canh cua da giac: 60
Nhap do day net ve (tu 1 den 5): 2
Dang khoi tao thuat toan xoay da giac de tao hoa van nghe thuat...
(Cửa sổ Turtle vẽ một đóa hoa màu tím đối xứng gồm 6 hình vuông cạnh 60 xoay 60 độ quanh tâm)""",
        explanations=[
            "Hàm ve_hoa_van_xoay nhận cấu hình linh hoạt theo số cạnh người dùng chọn.",
            "Vòng lặp ngoài chạy cố định 6 lần tạo 6 cánh hoa quanh trục; vòng lặp trong vẽ đa giác đều góc 360 / so_canh.",
            "Vẽ xong một đa giác, xoay phải 60 độ (360 / 6) để chuyển trục cho cánh hoa tiếp theo."
        ],
        tip_text="Học viên có thể thử nghiệm nhập so_canh = 3 (tam giác) hoặc so_canh = 5 (ngũ giác) để tạo ra các đóa hoa phong cách độc đáo khác."
    )


    # ==========================================================================
    # PHẦN 7: KHO 10 BÀI TOÁN THỰC HÀNH TỰ LUẬN ĐỀ THI
    # ==========================================================================
    add_chapter_divider(
        doc,
        "PHẦN 7: KHO 10 BÀI TOÁN THỰC HÀNH TỰ LUẬN ĐỀ THI (CHUẨN BỊ CHO BÀI THI 4 CÂU RANDOM - 40 PHÚT)",
        "Kho 10 bài toán lập trình hàm bám sát kiến thức khóa học, kèm code mẫu chuẩn và phân tích logic."
    )

    thuchanh_items = [
        ("BÀI TỰ LUẬN 1", "HÀM TÍNH TỔNG HAI SỐ A VÀ B CÓ KIỂM TRA ĐẦU VÀO", "def tinh_tong(a, b):\n    return a + b\n\na = float(input('Nhap so a: '))\nb = float(input('Nhap so b: '))\nprint('Tong hai so la:', tinh_tong(a, b))", "Nhap so a: 15\nNhap so b: 25\nTong hai so la: 40.0", ["Hàm nhận 2 tham số và trả về tổng a + b.", "Dùng float() để hỗ trợ tính cả số nguyên và số thực."]),
        ("BÀI TỰ LUẬN 2", "HÀM KIỂM TRA SỐ CHẴN HAY SỐ LẺ", "def kiem_tra_chan(n):\n    return n % 2 == 0\n\nn = int(input('Nhap so nguyen n: '))\nif kiem_tra_chan(n):\n    print(n, 'la so chan')\nelse:\n    print(n, 'la so le')", "Nhap so nguyen n: 8\n8 la so chan", ["Toán tử n % 2 == 0 trả về True nếu chia hết cho 2, ngược lại trả về False."]),
        ("BÀI TỰ LUẬN 3", "HÀM IN BẢNG CỬU CHƯƠNG TỪ 1 ĐẾN 10", "def in_bang_cuu_chuong(n):\n    print(f'Bang cuu chuong cua {n}:')\n    for i in range(1, 11):\n        print(f'{n} x {i} = {n * i}')\n\nn = int(input('Nhap so can in cuu chuong: '))\nin_bang_cuu_chuong(n)", "Nhap so can in cuu chuong: 7\nBang cuu chuong cua 7:\n7 x 1 = 7\n7 x 2 = 14\n...\n7 x 10 = 70", ["Dùng vòng lặp for i in range(1, 11) in 10 dòng phép nhân tương ứng."]),
        ("BÀI TỰ LUẬN 4", "HÀM TÍNH DIỆN TÍCH HÌNH TRÒN VỚI MATH.PI", "import math\n\ndef tinh_dien_tich_tron(r):\n    return math.pi * (r ** 2)\n\nr = float(input('Nhap ban kinh r: '))\nprint(f'Dien tich hinh tron la: {tinh_dien_tich_tron(r):.2f}')", "Nhap ban kinh r: 5\nDien tich hinh tron la: 78.54", ["Sử dụng hằng số math.pi và phép lũy thừa r ** 2.", "Định dạng :.2f để làm tròn 2 chữ số thập phân."]),
        ("BÀI TỰ LUẬN 5", "HÀM ĐẢO NGƯỢC XÂU KÝ TỰ (STRING SLICING)", "def dao_nguoc_chuoi(s):\n    return s[::-1]\n\ns = input('Nhap chuoi can dao: ')\nprint('Chuoi sau khi dao nguoc la:', dao_nguoc_chuoi(s))", "Nhap chuoi can dao: Python Sao Viet\nChuoi sau khi dao nguoc la: teiV oaS nohtyP", ["Cú pháp s[::-1] bước nhảy -1 giúp đảo ngược toàn bộ chuỗi ký tự tức thì."]),
        ("BÀI TỰ LUẬN 6", "HÀM TÍNH GIAI THỪA CỦA MỘT SỐ NGUYÊN N!", "def tinh_giai_thua(n):\n    if n == 0 or n == 1:\n        return 1\n    gt = 1\n    for i in range(2, n + 1):\n        gt = gt * i\n    return gt\n\nn = int(input('Nhap so n (n >= 0): '))\nprint(f'{n}! = {tinh_giai_thua(n)}')", "Nhap so n (n >= 0): 5\n5! = 120", ["Giai thừa n! = 1 * 2 * ... * n.", "Quy ước 0! = 1 và 1! = 1."]),
        ("BÀI TỰ LUẬN 7", "HÀM KIỂM TRA SỐ NGUYÊN TỐ TỐI ƯU", "def kiem_tra_nguyen_to(n):\n    if n < 2:\n        return False\n    for i in range(2, int(n ** 0.5) + 1):\n        if n % i == 0:\n            return False\n    return True\n\nn = int(input('Nhap so can kiem tra: '))\nif kiem_tra_nguyen_to(n):\n    print(n, 'la so nguyen to')\nelse:\n    print(n, 'khong phai so nguyen to')", "Nhap so can kiem tra: 17\n17 la so nguyen to", ["Số nguyên tố là số lớn hơn 1 và chỉ chia hết cho 1 và chính nó.", "Kiểm tra chia hết từ 2 đến căn bậc hai của n để tối ưu tốc độ."]),
        ("BÀI TỰ LUẬN 8", "HÀM TÍNH CHU VI & DIỆN TÍCH HÌNH CHỮ NHẬT", "def tinh_hcn(dai, rong):\n    chu_vi = (dai + rong) * 2\n    dien_tich = dai * rong\n    return chu_vi, dien_tich\n\ndai = float(input('Nhap chieu dai: '))\nrong = float(input('Nhap chieu rong: '))\ncv, dt = tinh_hcn(dai, rong)\nprint(f'Chu vi: {cv} | Dien tich: {dt}')", "Nhap chieu dai: 12.5\nNhap chieu rong: 7.5\nChu vi: 40.0 | Dien tich: 93.75", ["Hàm trả về đồng thời 2 giá trị (Tuple).", "Giải nén kết quả vào 2 biến cv và dt."]),
        ("BÀI TỰ LUẬN 9", "HÀM CHUYỂN ĐỔI ĐỘ C SANG ĐỘ F", "def c_sang_f(c):\n    return (c * 9 / 5) + 32\n\nc = float(input('Nhap nhiet do (do C): '))\nprint(f'{c} do C = {c_sang_f(c):.1f} do F')", "Nhap nhiet do (do C): 37\n37.0 do C = 98.6 do F", ["Công thức chuyển đổi chuẩn: F = (C * 9/5) + 32."]),
        ("BÀI TỰ LUẬN 10", "HÀM IN DANH SÁCH CÁC SỐ CHẴN TỪ 1 ĐẾN 100", "def in_so_chan_1_den_100():\n    print('Cac so chan tu 1 den 100:')\n    for i in range(2, 101, 2):\n        print(i, end=' ')\n    print()\n\nin_so_chan_1_den_100()", "Cac so chan tu 1 den 100:\n2 4 6 8 10 12 14 16 18 20 ... 98 100", ["Hàm range(2, 101, 2) bắt đầu từ 2, kết thúc ở 100 với bước nhảy 2.", "Tham số end=' ' giúp in các số nằm cùng trên một hàng ngang."])
    ]

    for item in thuchanh_items:
        add_lesson_box_bw(
            doc,
            badge=item[0],
            title=item[1],
            goal=f"Viết hàm giải quyết bài toán: {item[1]}.",
            metaphor="Khối xử lý logic độc lập, dễ dàng tái sử dụng nhiều lần trong chương trình.",
            code_text=item[2],
            output_text=item[3],
            explanations=item[4],
            tip_text="Hãy kiểm tra các trường hợp biên của dữ liệu đầu vào để code hoạt động ổn định nhất."
        )

    # ==========================================================================
    # BẢNG ĐÁNH GIÁ RÈN LUYỆN TOÀN KHÓA
    # ==========================================================================
    master_kpi_list = [
        "Chương 5 - Bài 12: Khởi động thư viện Turtle",
        "Chương 5 - Bài 13: Các lệnh điều khiển rùa cơ bản",
        "Chương 5 - Bài tập 1: Hình vuông 4 màu rực rỡ",
        "Chương 5 - Bài tập 2: Hình chữ nhật nét đậm màu tím",
        "Chương 5 - Bài tập 3: Vẽ 2 hình vuông xa nhau",
        "Chương 5 - Bài tập 4: Vẽ tam giác đều & đổi màu nền",
        "Chương 5 - Bài tập 5 (Mới): Ngôi sao 5 cánh viền đỏ ruột vàng",
        "Chương 5 - Bài tập 6 (Mới): Họa tiết xoắn ốc nghệ thuật",
        "Chương 1 - Bài 1: Cơ bản về chuỗi & cắt chuỗi slicing",
        "Chương 1 - Bài 2: Các phương thức xử lý chuỗi thường dùng",
        "Chương 1 - Bài 3: Định dạng chuỗi với f-string",
        "Chương 1 - Bài tập 1: Kiểm tra tên có số/khoảng trắng",
        "Chương 1 - Bài tập 2: Chuẩn hóa họ tên",
        "Chương 1 - Bài tập 3: Kiểm tra độ dài mật khẩu",
        "Chương 1 - Bài tập 4 (Mới): Đếm nguyên âm và phụ âm",
        "Chương 1 - Bài tập 5 (Mới): Ẩn thông tin số điện thoại",
        "Chương 2 - Bài 4: Kiểu dữ liệu List & duyệt mảng vòng lặp",
        "Chương 2 - Bài 5: Kiểu dữ liệu Dictionary (Khóa - Giá trị)",
        "Chương 2 - Bài tập 1: Hồ sơ học sinh bằng Dictionary",
        "Chương 2 - Bài tập 2: Tìm min/max & sắp xếp danh sách số",
        "Chương 2 - Bài tập 3: Quản lý giỏ đồ chơi (List lồng Dict)",
        "Chương 2 - Bài tập 4 (Mới): Tra cứu & cập nhật danh bạ SĐT",
        "Chương 2 - Bài tập 5 (Mới): Thống kê tần suất từ vựng",
        "Chương 3 - Bài 6: Khái niệm và cấu trúc hàm trong Python",
        "Chương 3 - Bài 7: Định nghĩa, tham số, return & None",
        "Chương 3 - Bài tập 1: Hàm kiểm tra chẵn/lẻ",
        "Chương 3 - Bài tập 2: Hàm tính tổng 1 đến n",
        "Chương 3 - Bài tập 3: Hàm in bảng cửu chương",
        "Chương 3 - Bài tập 4: Dự án máy tính 4 phép tính",
        "Chương 3 - Bài tập 5 (Mới): Hàm kiểm tra số đối xứng",
        "Chương 3 - Bài tập 6 (Mới): Hàm tính chỉ số BMI",
        "Chương 4 - Bài 8: Giới thiệu thư viện & cú pháp import",
        "Chương 4 - Bài 9: Thư viện Random sinh số & bốc thăm",
        "Chương 4 - Bài 10: Thư viện Math & các phép toán nâng cao",
        "Chương 4 - Bài 11: Thư viện Datetime làm việc với ngày giờ",
        "Chương 4 - Bài tập 1: Mô phỏng xúc xắc",
        "Chương 4 - Bài tập 2: Tính hình tròn, cầu, trụ",
        "Chương 4 - Bài tập 3: In ngày giờ hiện tại",
        "Chương 4 - Bài tập 4: Trò chơi đoán số 1-100",
        "Chương 4 - Bài tập 5 (Mới): Máy tạo mật khẩu ngẫu nhiên",
        "Chương 4 - Bài tập 6 (Mới): Đếm ngược ngày đến Tết",
        "Dự án 1: Chuỗi hình vuông tăng dần kích thước",
        "Dự án 2: Đa giác hoa văn màu sắc lồng nhau",
        "Dự án 3: Quản lý & phân loại điểm học sinh",
        "Dự án 4: Biểu đồ cột tăng trưởng Turtle",
        "Dự án 5: Bộ lọc an ninh quét mã CLB",
        "Dự án 6: Hoa văn xoay vòng từ đa giác đều",
        "Đề thi Thực hành: 10 Bài toán viết hàm hoàn chỉnh"
    ]
    add_kpi_table(doc, master_kpi_list)

    doc.save(DOCX_PATH)
    print(f"✅ Đã lưu Master Curriculum DOCX thành công: {DOCX_PATH}")

    print("📄 Đang chuyển đổi sang PDF...")
    try:
        convert(DOCX_PATH, PDF_PATH)
        print(f"🎉 Xuất Master Curriculum PDF thành công: {PDF_PATH}")
    except Exception as e:
        print(f"❌ Lỗi khi convert PDF: {e}")

if __name__ == "__main__":
    build_master_doc()
