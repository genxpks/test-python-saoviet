"""
generate_quiz_bank_doc.py
Biên soạn ngân hàng 120 câu hỏi trắc nghiệm (6 dạng) & 10 bài toán tự luận thực hành
Đơn vị: TIN HỌC SAO VIỆT THỦ ĐỨC
Xuất bản: DOCX, PDF, JSON và JavaScript module.
"""

import os
import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx2pdf import convert

sys.stdout.reconfigure(encoding='utf-8')

CUR_DIR = os.path.dirname(os.path.abspath(__file__))
PKG_DIR = os.path.dirname(CUR_DIR)
BANK_DIR = os.path.join(PKG_DIR, "Ngan_Hang_120_Cau_Hoi")
WEB_DIR = os.path.join(PKG_DIR, "Web_On_Thi_Trac_Nghiem")
os.makedirs(BANK_DIR, exist_ok=True)
os.makedirs(WEB_DIR, exist_ok=True)

DOCX_PATH = os.path.join(BANK_DIR, "Ngan_Hang_120_Cau_Trac_Nghiem_Va_10_Bai_Thuc_Hanh.docx")
PDF_PATH  = os.path.join(BANK_DIR, "Ngan_Hang_120_Cau_Trac_Nghiem_Va_10_Bai_Thuc_Hanh.pdf")
JSON_PATH = os.path.join(BANK_DIR, "questions_bank_full.json")
JS_PATH   = os.path.join(WEB_DIR, "questions.js")

BLACK     = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(60, 60, 60)
WHITE     = RGBColor(255, 255, 255)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_border(cell, top="single", bottom="single", left="single", right="single", sz="8", color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tcBorders>''')
    tcPr.append(tcBorders)

# Import and build the rich 120 question bank
def get_all_120_questions():
    questions = []
    
    # ----------------------------------------------------
    # DẠNG 1: TRẮC NGHIỆM 4 LỰA CHỌN ABCD (40 CÂU: ID 1 -> 40)
    # ----------------------------------------------------
    mc_data = [
        ("Trong Python, cú pháp nào sau đây dùng để khai báo một hàm mới?", ["function my_func():", "def my_func():", "create my_func():", "func my_func():"], 1, "Từ khóa 'def' là từ khóa chuẩn trong Python dùng để định nghĩa một hàm mới."),
        ("Kết quả của đoạn lệnh: s = 'Python'; print(s[0]) là gì?", ["P", "y", "n", "Báo lỗi"], 0, "Chỉ số index trong Python bắt đầu từ vị trí 0, do đó s[0] lấy ký tự đầu tiên là 'P'."),
        ("Để lấy ký tự cuối cùng của chuỗi s, ta dùng chỉ số nào?", ["s[1]", "s[end]", "s[-1]", "s[last]"], 2, "Chỉ số âm -1 đại diện cho phần tử cuối cùng của chuỗi hoặc danh sách trong Python."),
        ("Phương thức nào dùng để viết hoa chữ cái đầu tiên của từng từ trong chuỗi?", ["upper()", "capitalize()", "title()", "lower()"], 2, "Phương thức .title() tự động viết hoa chữ cái đầu của mỗi từ trong chuỗi."),
        ("Phương thức nào dùng để cắt bỏ khoảng trắng dư thừa ở hai đầu chuỗi?", ["clean()", "strip()", "trim()", "cut()"], 1, "Phương thức .strip() loại bỏ toàn bộ khoảng trắng ở đầu và cuối chuỗi."),
        ("Kết quả của đoạn lệnh: len('Sao Viet') là bao nhiêu?", ["7", "8", "9", "6"], 1, "Chuỗi 'Sao Viet' có 8 ký tự bao gồm 7 chữ cái và 1 khoảng trắng ở giữa."),
        ("Để kiểm tra một chuỗi có toàn bộ là chữ số hay không, ta dùng hàm nào?", ["s.isnumber()", "s.isdigit()", "s.isnumeric_only()", "s.check_digit()"], 1, "Phương thức .isdigit() trả về True nếu tất cả ký tự trong chuỗi là chữ số từ 0-9."),
        ("Kết quả in của f-string: diem = 8.666; print(f'{diem:.2f}') là gì?", ["8.6", "8.66", "8.67", "8.666"], 2, "Cú pháp {diem:.2f} làm tròn số thực đến 2 chữ số thập phân, 8.666 làm tròn lên thành 8.67."),
        ("Trong Python, danh sách List được khai báo bằng cặp ngoặc nào?", ["( )", "{ }", "[ ]", "< >"], 2, "Danh sách List trong Python được định nghĩa bằng cặp ngoặc vuông [ ]."),
        ("Để thêm một phần tử vào cuối danh sách a, ta dùng phương thức nào?", ["a.add(x)", "a.append(x)", "a.insert_last(x)", "a.push(x)"], 1, "Phương thức .append(x) thêm phần tử x vào vị trí cuối cùng của List."),
        ("Kết quả của đoạn lệnh: a = [10, 20, 30]; print(a[-1]) là gì?", ["10", "20", "30", "Báo lỗi"], 2, "a[-1] truy xuất phần tử cuối cùng của danh sách a, chính là số 30."),
        ("Để tìm giá trị lớn nhất trong danh sách số a, ta dùng hàm nào?", ["max(a)", "a.maximum()", "top(a)", "a.largest()"], 0, "Hàm built-in max(a) tự động tìm và trả về phần tử có giá trị lớn nhất."),
        ("Để sắp xếp danh sách a tăng dần trực tiếp, ta dùng lệnh nào?", ["a.order()", "a.sort()", "sort(a)", "a.arrange()"], 1, "Phương thức a.sort() sắp xếp các phần tử của List theo thứ tự tăng dần tại chỗ."),
        ("Dictionary trong Python lưu trữ dữ liệu dưới dạng cấu trúc nào?", ["Chỉ gồm các giá trị số", "Cặp Khóa — Giá trị (Key — Value)", "Mảng 2 chiều cố định", "Hàng đợi FIFO"], 1, "Dictionary lưu trữ dữ liệu dưới dạng các cặp Khóa - Giá trị {key: value}."),
        ("Để lấy tất cả các cặp (Key, Value) từ từ điển d để duyệt vòng lặp, ta dùng:", ["d.pairs()", "d.items()", "d.all()", "d.elements()"], 1, "Phương thức d.items() trả về danh sách các tuple chứa cả khóa và giá trị."),
        ("Lệnh nào dùng để kết thúc hàm và trả kết quả về nơi gọi?", ["stop", "exit", "return", "break"], 2, "Từ khóa return kết thúc việc thực thi hàm và trả về giá trị tính toán được."),
        ("Nếu một hàm không có lệnh return, giá trị trả về mặc định của hàm là gì?", ["0", "False", "None", "'' (Chuỗi rỗng)"], 2, "Trong Python, hàm không có return hoặc chỉ gọi return không giá trị sẽ trả về None."),
        ("Để nạp thư viện toán học Math vào chương trình, cú pháp chuẩn là gì?", ["include math", "import math", "using math", "require math"], 1, "Từ khóa 'import' dùng để nạp các module/thư viện trong Python."),
        ("Hàm nào trong thư viện random dùng để sinh một số nguyên ngẫu nhiên từ a đến b?", ["random.rand(a, b)", "random.randint(a, b)", "random.choice(a, b)", "random.integer(a, b)"], 1, "random.randint(a, b) sinh số nguyên ngẫu nhiên trong đoạn [a, b] (bao gồm cả a và b)."),
        ("Để bốc thăm ngẫu nhiên một phần tử từ một danh sách list, ta dùng:", ["random.randint(list)", "random.pick(list)", "random.choice(list)", "random.select(list)"], 2, "Hàm random.choice(danh_sach) chọn ngẫu nhiên 1 phần tử từ tập hợp."),
        ("Hằng số số Pi trong thư viện math được truy xuất bằng cú pháp nào?", ["math.PI()", "math.pi", "math.PI_VALUE", "math.get_pi()"], 1, "math.pi là biến hằng số lưu giá trị xấp xỉ 3.141592653589793."),
        ("Kết quả của math.sqrt(25) là kiểu dữ liệu gì và giá trị bao nhiêu?", ["Số nguyên 5", "Số thực 5.0", "Chuỗi '5'", "Số phức 5j"], 1, "Hàm math.sqrt() luôn luôn trả về kết quả dưới dạng số thực (float) là 5.0."),
        ("Để tính giai thừa của số 5 (5! = 120), ta gọi hàm nào?", ["math.fact(5)", "math.factorial(5)", "math.pow(5)", "math.giai_thua(5)"], 1, "Hàm math.factorial(n) tính toán giai thừa của số nguyên n không âm."),
        ("Để lấy thời gian thực hiện tại của hệ thống, ta dùng lệnh nào?", ["datetime.now()", "datetime.datetime.now()", "datetime.get_current_time()", "time.current()"], 1, "datetime.datetime.now() trả về đối tượng ngày giờ hiện tại đầy đủ."),
        ("Trong thư viện Turtle, lệnh nào dùng để đưa chú rùa tiến thẳng về phía trước 100 bước?", ["but_ve.move(100)", "but_ve.forward(100)", "but_ve.go(100)", "but_ve.step(100)"], 1, "Lệnh but_ve.forward(100) hoặc but_ve.fd(100) di chuyển rùa tiến thẳng."),
        ("Lệnh nào dùng để xoay góc chú rùa sang bên phải 90 độ?", ["but_ve.turn_right(90)", "but_ve.right(90)", "but_ve.rotate_right(90)", "but_ve.r(90)"], 1, "Lệnh but_ve.right(90) hoặc but_ve.rt(90) xoay rùa sang phải 90 độ."),
        ("Để nhấc bút vẽ lên không để lại nét mực khi rùa di chuyển, ta dùng lệnh:", ["but_ve.penoff()", "but_ve.penup()", "but_ve.lift()", "but_ve.hide_pen()"], 1, "but_ve.penup() hoặc but_ve.up() nhấc bút lên khỏi mặt phẳng vẽ."),
        ("Để đặt độ dày của nét vẽ Turtle là 4 điểm ảnh, ta dùng lệnh nào?", ["but_ve.width(4)", "but_ve.pensize(4)", "but_ve.thickness(4)", "Cả A và B đều đúng"], 3, "Trong Turtle, cả pensize(4) và width(4) đều dùng để thiết lập độ dày nét vẽ."),
        ("Lệnh nào giữ cho cửa sổ đồ họa Turtle luôn hiển thị sau khi vẽ xong?", ["turtle.stay()", "turtle.keep()", "turtle.done()", "turtle.stop()"], 2, "turtle.done() hoặc turtle.mainloop() giữ cửa sổ đồ họa mở để người dùng quan sát."),
        ("Để đổi màu nền của cửa sổ Turtle thành màu xanh lá, ta dùng lệnh:", ["screen.bgcolor('green')", "turtle.background('green')", "screen.color('green')", "turtle.set_screen('green')"], 0, "screen.bgcolor('green') thiết lập màu nền (background color) cho màn hình vẽ."),
        ("Kết quả của đoạn lệnh: print('10' + '20') là gì?", ["30", "1020", "Báo lỗi", "10 20"], 1, "Dấu cộng giữa 2 chuỗi là phép ghép nối chuỗi (concatenation), ghép '10' và '20' thành '1020'."),
        ("Lệnh nào dùng để chuyển chuỗi '123' thành số nguyên 123?", ["str(123)", "int('123')", "float('123')", "number('123')"], 1, "Hàm int() thực hiện ép kiểu từ chuỗi ký tự sang số nguyên."),
        ("Kết quả của phép chia lấy dư: 17 % 5 là bao nhiêu?", ["3", "2", "3.4", "1"], 1, "17 chia 5 được 3 dư 2, toán tử % lấy phần dư nên kết quả là 2."),
        ("Trong vòng lặp for i in range(1, 5), biến i sẽ nhận lần lượt các giá trị nào?", ["1, 2, 3, 4, 5", "1, 2, 3, 4", "0, 1, 2, 3, 4", "0, 1, 2, 3, 4, 5"], 1, "Hàm range(start, stop) chạy từ start đến stop - 1, do đó range(1, 5) gồm 1, 2, 3, 4."),
        ("Lệnh nào dùng để thoát khỏi vòng lặp ngay lập tức?", ["continue", "exit", "break", "return"], 2, "Từ khóa break ngắt và thoát khỏi vòng lặp gần nhất ngay lập tức."),
        ("Cú pháp s[::-1] trên một chuỗi s có tác dụng gì?", ["Lấy ký tự đầu", "Lấy ký tự cuối", "Đảo ngược chuỗi", "Xóa chuỗi"], 2, "Slicing bước nhảy -1 (s[::-1]) duyệt chuỗi từ cuối lên đầu, tạo ra chuỗi đảo ngược."),
        ("Để kiểm tra một chuỗi chỉ gồm chữ cái và số (không chứa khoảng trắng hay ký tự đặc biệt), ta dùng:", ["s.isalpha()", "s.isdigit()", "s.isalnum()", "s.isspace()"], 2, "s.isalnum() (is alpha-numeric) kiểm tra chuỗi chỉ gồm chữ cái và chữ số."),
        ("Để tô màu kín cho một hình vẽ khép kín trong Turtle, ta kẹp các lệnh vẽ ở giữa cặp lệnh nào?", ["begin_fill() và end_fill()", "start_color() và stop_color()", "fill_on() và fill_off()", "paint_begin() và paint_end()"], 0, "Cặp lệnh but_ve.begin_fill() và but_ve.end_fill() tự động tô màu kín hình đa giác."),
        ("Tổng các góc ngoài của một hình ngũ giác đều là bao nhiêu độ?", ["180 độ", "360 độ", "540 độ", "720 độ"], 1, "Tổng góc ngoài của mọi đa giác lồi luôn luôn bằng 360 độ (mỗi góc xoay là 360 / 5 = 72 độ)."),
        ("Đoạn mã: a = [1, 2, 3]; a.append([4, 5]); len(a) cho kết quả là:", ["5", "4", "3", "Báo lỗi"], 1, "Phương thức .append([4, 5]) thêm cả danh sách con [4, 5] như 1 phần tử duy nhất, len(a) là 4.")
    ]

    q_id = 1
    for item in mc_data:
        questions.append({
            "id": q_id,
            "type": "single_choice",
            "type_name": "Trắc nghiệm ABCD (1 đáp án)",
            "question": item[0],
            "options": item[1],
            "correct_answer": item[2],
            "explanation": item[3]
        })
        q_id += 1

    # ----------------------------------------------------
    # DẠNG 2: ĐÚNG / SAI (TRUE / FALSE) (20 CÂU: ID 41 -> 60)
    # ----------------------------------------------------
    tf_data = [
        ("Trong Python, chuỗi ký tự (String) là kiểu dữ liệu có thể thay đổi (Mutable) trực tiếp từng ký tự bằng phép gán s[0] = 'X'.", False, "Sai. Chuỗi trong Python là kiểu bất biến (Immutable), không thể thay đổi giá trị từng ký tự trực tiếp."),
        ("Hàm random.randint(1, 6) có thể sinh ra số 6.", True, "Đúng. random.randint(a, b) bao gồm cả giá trị cận trên b (1 đến 6)."),
        ("Từ khóa return trong hàm Python sẽ kết thúc hàm ngay lập tức khi được thực thi.", True, "Đúng. Khi gặp lệnh return, luồng chương trình thoát khỏi hàm và trả về giá trị."),
        ("Trong Dictionary của Python, hai khóa (Key) khác nhau có thể chứa hai giá trị (Value) giống nhau.", True, "Đúng. Các Key phải là duy nhất, nhưng các Value gắn với Key thì hoàn toàn có thể trùng nhau."),
        ("Góc quay ngoài của tam giác đều khi vẽ bằng thư viện Turtle là 60 độ.", False, "Sai. Góc ngoài của tam giác đều là 120 độ (360 / 3 = 120 độ). 60 độ là góc trong."),
        ("Phương thức s.strip() sẽ xóa toàn bộ khoảng trắng ở giữa các từ trong câu.", False, "Sai. .strip() chỉ xóa khoảng trắng thừa ở 2 đầu chuỗi, không xóa khoảng trắng ở giữa các từ."),
        ("Toán tử len() có thể dùng để đếm số phần tử của cả String, List và Dictionary.", True, "Đúng. Hàm len() áp dụng được cho mọi cấu trúc tuần tự và tập hợp trong Python."),
        ("Phương thức a.sort() trả về một danh sách mới mà không làm thay đổi danh sách ban đầu.", False, "Sai. a.sort() sắp xếp trực tiếp trên danh sách gốc và trả về None. Muốn tạo danh sách mới ta dùng hàm sorted(a)."),
        ("Để import hàm sqrt từ thư viện math, ta có thể viết: from math import sqrt.", True, "Đúng. Đây là cú pháp import trực tiếp hàm từ module."),
        ("Lệnh but_ve.penup() sẽ xóa toàn bộ các nét vẽ đã vẽ trước đó trên màn hình.", False, "Sai. but_ve.penup() chỉ nhấc bút không vẽ tiếp, muốn xóa màn hình ta dùng lệnh but_ve.clear() hoặc reset()."),
        ("Tên biến trong Python có thể bắt đầu bằng chữ số (ví dụ: 1ten = 'An').", False, "Sai. Tên biến trong Python không được phép bắt đầu bằng chữ số."),
        ("Khối lệnh con trong hàm Python bắt buộc phải được thụt lề (Indentation) đồng nhất.", True, "Đúng. Python dùng khoảng thụt lề (thường là 4 dấu cách hoặc 1 tab) để xác định khối lệnh."),
        ("Hàm math.pow(2, 3) trả về kết quả là số thực 8.0.", True, "Đúng. Hàm math.pow() luôn luôn trả về kiểu số thực float."),
        ("Khi dùng vòng lặp for i in range(5), biến i sẽ bắt đầu từ 1 đến 5.", False, "Sai. range(5) bắt đầu từ 0 đến 4."),
        ("Dictionary trong Python cho phép truy xuất phần tử theo số thứ tự index [0], [1].", False, "Sai. Dictionary truy xuất qua Khóa Key (ví dụ d['ten']), không truy xuất qua vị trí index."),
        ("Lệnh but_ve.circle(50) trong Turtle vẽ hình tròn có đường kính là 50.", False, "Sai. Tham số 50 là Bán kính (radius) của hình tròn, đường kính sẽ là 100."),
        ("Trong f-string, biểu thức bên trong dấu ngoặc nhọn {} có thể là một phép tính toán học (ví dụ: f'{2 + 3}').", True, "Đúng. f-string cho phép tính toán trực tiếp các biểu thức nằm trong ngoặc nhọn."),
        ("Lệnh datetime.datetime.now().year trả về năm hiện tại dưới dạng số nguyên.", True, "Đúng. Thuộc tính .year trả về năm dạng int (ví dụ: 2026)."),
        ("Toán tử 'in' có thể dùng để kiểm tra xem một phần tử có nằm trong List hay không.", True, "Đúng. Cú pháp 'x in a' trả về True nếu x tồn tại trong danh sách a."),
        ("Lệnh turtle.speed(0) đặt tốc độ vẽ của chú rùa là chậm nhất.", False, "Sai. speed(0) là tốc độ vẽ nhanh nhất (không có độ trễ hoạt ảnh).")
    ]

    for item in tf_data:
        questions.append({
            "id": q_id,
            "type": "true_false",
            "type_name": "Trắc nghiệm Đúng / Sai",
            "question": item[0],
            "options": ["Đúng (True)", "Sai (False)"],
            "correct_answer": 0 if item[1] else 1,
            "explanation": item[2]
        })
        q_id += 1

    # ----------------------------------------------------
    # DẠNG 3: CHỌN NHIỀU ĐÁP ÁN ĐÚNG (CHECKBOX) (20 CÂU: ID 61 -> 80)
    # ----------------------------------------------------
    multi_data = [
        ("Những phương thức nào sau đây KHÔNG làm thay đổi chuỗi gốc ban đầu mà trả về chuỗi mới?", ["upper()", "lower()", "strip()", "title()"], [0, 1, 2, 3], "Tất cả các phương thức xử lý chuỗi trong Python đều không sửa chuỗi gốc vì String là bất biến."),
        ("Những cú pháp nào sau đây là cách import thư viện hợp lệ trong Python?", ["import math", "from math import pi, sqrt", "import datetime as dt", "using math"], [0, 1, 2], "Các đáp án A, B, C đều là cú pháp import chuẩn của Python. 'using' là cú pháp của C#/C++."),
        ("Những lệnh nào sau đây dùng để điều khiển chuyển động di chuyển của chú rùa Turtle?", ["forward(100)", "backward(50)", "right(90)", "circle(40)"], [0, 1, 2, 3], "Cả 4 lệnh trên đều trực tiếp điều khiển hướng đi và vị trí của con trỏ bút vẽ Turtle."),
        ("Các kiểu dữ liệu nào sau đây lưu trữ dữ liệu theo thứ tự (Ordered Sequence) và có thể cắt chuỗi/slicing?", ["String (Chuỗi)", "List (Danh sách)", "Tuple", "Dictionary (trong các bản Python cũ)"], [0, 1, 2], "String, List và Tuple là các kiểu dữ liệu tuần tự có chỉ mục vị trí index rõ ràng."),
        ("Những hàm nào sau đây thuộc thư viện toán học math có sẵn trong Python?", ["sqrt()", "factorial()", "pow()", "randint()"], [0, 1, 2], "sqrt, factorial, pow thuộc math. Còn randint thuộc thư viện random."),
        ("Những câu lệnh nào sau đây giúp tạo một danh sách rỗng trong Python?", ["a = []", "a = list()", "a = {}", "a = ()"], [0, 1], "a = [] và a = list() tạo danh sách rỗng. {} tạo dictionary rỗng, () tạo tuple rỗng."),
        ("Những thao tác nào sau đây có thể thực hiện trên một List trong Python?", ["append()", "remove()", "sort()", "pop()"], [0, 1, 2, 3], "Tất cả 4 phương thức trên đều là các thao tác chuẩn có sẵn trên kiểu dữ liệu List."),
        ("Những ký tự nào sau đây là ký tự điều khiển đặc biệt (Escape sequence) trong chuỗi Python?", ["\\n (xuống dòng)", "\\t (thụt lề tab)", "\\\\ (in dấu gạch chéo)", "\\p (in hoa)"], [0, 1, 2], "\\n, \\t, \\\\ là các escape sequence chuẩn trong Python."),
        ("Các cách nào sau đây giúp duyệt qua các phần tử của một danh sách a = [10, 20, 30]?", ["for x in a:", "for i in range(len(a)):", "for x in a.items():", "while i < len(a):"], [0, 1, 3], "Có thể duyệt trực tiếp qua giá trị, qua chỉ số index với range(len(a)) hoặc dùng while loop."),
        ("Những câu lệnh nào sau đây dùng để thiết lập màu sắc trong đồ họa Turtle?", ["color('red')", "color('blue', 'yellow')", "pencolor('green')", "fillcolor('orange')"], [0, 1, 2, 3], "Turtle hỗ trợ đầy đủ các lệnh đặt màu viền (pencolor), màu tô (fillcolor) và đặt cùng lúc (color)."),
        ("Để kiểm tra một số n có phải là số chẵn dương, những điều kiện nào sau đây là đúng?", ["n % 2 == 0 and n > 0", "n > 0 and n % 2 == 0", "n % 2 != 1 and n >= 2", "n % 2 == 1 and n > 0"], [0, 1, 2], "Các phương án A, B, C đều biểu diễn chính xác điều kiện số nguyên chẵn dương."),
        ("Những thuộc tính nào sau đây có thể lấy từ đối tượng now = datetime.datetime.now()?", ["now.year", "now.month", "now.hour", "now.minute"], [0, 1, 2, 3], "Đối tượng datetime chứa đầy đủ các thuộc tính year, month, day, hour, minute, second."),
        ("Những cách nào sau đây dùng để xóa một phần tử khỏi danh sách a?", ["a.remove(x)", "del a[0]", "a.pop()", "a.delete(x)"], [0, 1, 2], "remove(x), del và pop() là 3 cách xóa phần tử hợp lệ trong Python."),
        ("Những nhận định nào sau đây là ĐÚNG về hàm (function) trong Python?", ["Giúp tái sử dụng mã nguồn", "Khai báo bắt đầu bằng từ khóa def", "Có thể nhận tham số đầu vào", "Bắt buộc phải có lệnh print"], [0, 1, 2], "Hàm giúp tái sử dụng code, dùng def và nhận tham số. Hàm không bắt buộc phải có print."),
        ("Những giá trị nào sau đây khi chuyển sang kiểu Boolean bool(x) sẽ cho kết quả là False?", ["0", "'' (Chuỗi rỗng)", "[] (List rỗng)", "'0' (Chuỗi số 0)"], [0, 1, 2], "Số 0, chuỗi rỗng và list rỗng mang giá trị Falsy. Chuỗi '0' có độ dài 1 nên mang giá trị True."),
        ("Những phương thức nào sau đây giúp tìm kiếm vị trí của một chuỗi con?", ["find()", "index()", "search()", "locate()"], [0, 1], "find() và index() là 2 phương thức tìm kiếm vị trí chuỗi con chuẩn của kiểu String."),
        ("Những tham số nào có thể truyền vào hàm print() để tùy biến hiển thị?", ["sep=' '", "end='\\n'", "file=sys.stdout", "style='bold'"], [0, 1, 2], "sep và end là 2 tham số quan trọng nhất thường dùng trong print()."),
        ("Những hình nào sau đây có thể vẽ dễ dàng bằng vòng lặp trong Turtle?", ["Hình tam giác đều", "Hình vuông", "Hình ngôi sao 5 cánh", "Hình xoắn ốc"], [0, 1, 2, 3], "Tất cả các hình học đối xứng đều vẽ rất đẹp mắt bằng vòng lặp for trong Turtle."),
        ("Những hàm nào sau đây của module random trả về kết quả ngẫu nhiên?", ["random.randint(1, 10)", "random.choice(['A', 'B'])", "random.random()", "random.shuffle(list)"], [0, 1, 2, 3], "Tất cả các hàm trên đều thuộc thư viện random dùng để thao tác ngẫu nhiên hóa."),
        ("Những phép toán nào sau đây cho kết quả là số thực (float)?", ["10 / 2", "math.sqrt(9)", "2.5 * 2", "10 // 2"], [0, 1, 2], "Phép chia /, hàm sqrt() và phép tính với số thực đều trả về float. Phép chia nguyên // trả về int.")
    ]

    for item in multi_data:
        questions.append({
            "id": q_id,
            "type": "multiple_choice",
            "type_name": "Trắc nghiệm chọn nhiều đáp án đúng",
            "question": item[0],
            "options": item[1],
            "correct_answer": item[2],
            "explanation": item[3]
        })
        q_id += 1

    # ----------------------------------------------------
    # DẠNG 4: ĐIỀN VÀO CHỖ TRỐNG (FILL IN THE BLANK) (15 CÂU: ID 81 -> 95)
    # ----------------------------------------------------
    fill_data = [
        ("Từ khóa để bắt đầu định nghĩa một hàm trong Python là từ khóa: ___.", "def", "Cú pháp chuẩn: def ten_ham(tham_so):"),
        ("Để nhấc bút vẽ lên trong Turtle mà không vẽ nét khi di chuyển, ta dùng lệnh: but_ve.___().", "penup", "but_ve.penup() nhấc bút lên khỏi mặt vẽ."),
        ("Trong f-string, để định dạng số thực x làm tròn 2 chữ số thập phân ta viết: f'{x:.___}'.", ".2f", "Ký hiệu .2f có nghĩa là float được làm tròn 2 chữ số sau dấu thập phân."),
        ("Để sinh số nguyên ngẫu nhiên trong khoảng từ 1 đến 100, ta dùng hàm: random.___(1, 100).", "randint", "Hàm random.randint(a, b) sinh số nguyên ngẫu nhiên."),
        ("Để thêm phần tử mới vào cuối danh sách List a, ta gọi phương thức: a.___(gia_tri).", "append", "Phương thức .append() thêm phần tử vào đuôi danh sách."),
        ("Để tính căn bậc hai của số 64 trong thư viện math, ta viết: math.___(64).", "sqrt", "Hàm math.sqrt() viết tắt của Square Root (căn bậc hai)."),
        ("Để xoay chú rùa sang bên phải 90 độ, ta dùng lệnh: but_ve.___(90).", "right", "but_ve.right(90) điều khiển rùa quay sang phải."),
        ("Phương thức dùng để loại bỏ khoảng trắng thừa ở hai đầu chuỗi là: s.___().", "strip", "Phương thức .strip() cắt bỏ khoảng trắng hai đầu."),
        ("Để lấy tổng số ký tự của một chuỗi hoặc số phần tử của danh sách, ta dùng hàm: ___(a).", "len", "Hàm built-in len() trả về độ dài (length) của đối tượng."),
        ("Để kết thúc vòng lặp ngay lập tức khi thỏa mãn điều kiện, ta dùng lệnh: ___.", "break", "Từ khóa break ngắt vòng lặp ngay lập tức."),
        ("Để lấy thời gian hiện tại trong module datetime, ta gọi: datetime.datetime.___().", "now", "Phương thức now() lấy mốc thời gian thực hiện tại."),
        ("Hằng số số Pi trong thư viện math được viết là: math.___.", "pi", "Hằng số math.pi lưu giá trị 3.141592653589793."),
        ("Để duyệt qua tất cả các cặp khóa - giá trị của từ điển d, ta gọi: for k, v in d.___():", "items", "Phương thức d.items() trả về danh sách các cặp (key, value)."),
        ("Để đảo ngược toàn bộ chuỗi s bằng kỹ thuật Slicing, ta viết: s[:::___].", "-1", "Bước nhảy -1 trong cú pháp s[::-1] đảo ngược chuỗi."),
        ("Để giữ cửa sổ đồ họa Turtle luôn hiển thị sau khi chạy xong, ta gọi lệnh: turtle.___().", "done", "Lệnh turtle.done() giữ cửa sổ màn hình đồ họa mở.")
    ]

    for item in fill_data:
        questions.append({
            "id": q_id,
            "type": "fill_blank",
            "type_name": "Điền vào chỗ trống",
            "question": item[0],
            "correct_answer": item[1],
            "explanation": item[2]
        })
        q_id += 1

    # ----------------------------------------------------
    # DẠNG 5: SẮP XẾP THỨ TỰ CÁC BƯỚC / DÒNG CODE (15 CÂU: ID 96 -> 110)
    # ----------------------------------------------------
    order_data = [
        ("Hãy sắp xếp các dòng lệnh sau theo đúng thứ tự để tạo và in thông tin học sinh bằng Dictionary:", ["hoc_sinh = {}", "hoc_sinh['ten'] = 'Minh'", "hoc_sinh['diem'] = 9.5", "print(hoc_sinh)"], [0, 1, 2, 3], "Quy trình: Tạo dict rỗng -> Gán tên -> Gán điểm -> In kết quả."),
        ("Hãy sắp xếp các bước để vẽ một hình vuông 4 cạnh trong Turtle:", ["import turtle", "but_ve = turtle.Turtle()", "for i in range(4):", "    but_ve.forward(100)", "    but_ve.right(90)", "turtle.done()"], [0, 1, 2, 3, 4, 5], "Quy trình: Import thư viện -> Tạo bút vẽ -> Lặp 4 lần -> Tiến 100 -> Quay 90 -> Giữ màn hình."),
        ("Hãy sắp xếp các dòng lệnh để tính tổng các số từ 1 đến 5 bằng vòng lặp for:", ["tong = 0", "for i in range(1, 6):", "    tong = tong + i", "print('Tong la:', tong)"], [0, 1, 2, 3], "Quy trình: Khởi tạo biến tổng = 0 -> Duyệt i từ 1 đến 5 -> Cộng dồn i vào tổng -> In tổng."),
        ("Hãy sắp xếp các bước chuẩn để định nghĩa và gọi một hàm tính diện tích hình chữ nhật:", ["def tinh_dien_tich(dai, rong):", "    return dai * rong", "ket_qua = tinh_dien_tich(5, 4)", "print('Dien tich la:', ket_qua)"], [0, 1, 2, 3], "Quy trình: Khai báo hàm def -> Viết thân hàm return -> Gọi hàm truyền đối số -> In kết quả."),
        ("Hãy sắp xếp các dòng lệnh để bốc thăm ngẫu nhiên một món quà từ danh sách:", ["import random", "qua_tang = ['But', 'Sach', 'Cap']", "mon_qua = random.choice(qua_tang)", "print('Mon qua trung thuong la:', mon_qua)"], [0, 1, 2, 3], "Quy trình: Import module random -> Tạo danh sách quà -> Gọi random.choice() -> In kết quả."),
        ("Hãy sắp xếp các bước để nhập một số nguyên và kiểm tra số chẵn lẻ:", ["n = int(input('Nhap n: '))", "if n % 2 == 0:", "    print(n, 'la so chan')", "else:", "    print(n, 'la so le')"], [0, 1, 2, 3, 4], "Quy trình: Nhập dữ liệu ép kiểu int -> Kiểm tra điều kiện n % 2 == 0 -> In kết quả chẵn -> In kết quả lẻ."),
        ("Hãy sắp xếp các dòng lệnh để nhập 3 số vào List và sắp xếp tăng dần:", ["ds = []", "for i in range(3):", "    so = int(input('Nhap so: '))", "    ds.append(so)", "ds.sort()", "print(ds)"], [0, 1, 2, 3, 4, 5], "Quy trình: Tạo list rỗng -> Lặp 3 lần nhập số và append -> Sắp xếp bằng sort() -> In danh sách."),
        ("Hãy sắp xếp các bước để chuẩn hóa một chuỗi họ tên người dùng:", ["s = '   nguyen van an   '", "s = s.strip()", "s = s.title()", "print('Ten chuan hoa:', s)"], [0, 1, 2, 3], "Quy trình: Gán chuỗi thô -> Cắt khoảng trắng thừa .strip() -> Viết hoa chữ cái đầu .title() -> In kết quả."),
        ("Hãy sắp xếp các bước để vẽ tam giác đều tô màu đỏ trong Turtle:", ["but_ve.color('red', 'red')", "but_ve.begin_fill()", "for i in range(3):", "    but_ve.forward(100)", "    but_ve.left(120)", "but_ve.end_fill()"], [0, 1, 2, 3, 4, 5], "Quy trình: Đặt màu bút và màu tô -> begin_fill() -> Lặp 3 lần vẽ tam giác -> end_fill() kết thúc tô màu."),
        ("Hãy sắp xếp các dòng lệnh để đếm số lượng ký tự số trong một chuỗi:", ["s = 'Python2026'", "dem = 0", "for ch in s:", "    if ch.isdigit():", "        dem += 1", "print('So chu so la:', dem)"], [0, 1, 2, 3, 4, 5], "Quy trình: Khởi tạo chuỗi và biến đếm -> Duyệt từng ký tự -> Kiểm tra .isdigit() -> Tăng đếm -> In kết quả."),
        ("Hãy sắp xếp các bước để tính căn bậc hai của một số nhập từ bàn phím:", ["import math", "x = float(input('Nhap so: '))", "can_bac_hai = math.sqrt(x)", "print(f'Can bac hai la: {can_bac_hai:.2f}')"], [0, 1, 2, 3], "Quy trình: Import math -> Nhập số float -> Tính math.sqrt() -> In định dạng .2f."),
        ("Hãy sắp xếp các bước để in bảng cửu chương 5 từ 1 đến 10:", ["n = 5", "print(f'Bang cuu chuong {n}:')", "for i in range(1, 11):", "    print(f'{n} x {i} = {n * i}')"], [0, 1, 2, 3], "Quy trình: Gán n = 5 -> In tiêu đề -> Vòng lặp for 1..10 -> In từng dòng phép tính nhân."),
        ("Hãy sắp xếp các bước để lưu danh bạ điện thoại và tra cứu theo tên:", ["danh_ba = {'An': '090123'}", "ten = input('Nhap ten can tra: ')", "if ten in danh_ba:", "    print('SDT:', danh_ba[ten])"], [0, 1, 2, 3], "Quy trình: Tạo từ điển danh bạ -> Nhập tên tra cứu -> Kiểm tra 'in danh_ba' -> In số điện thoại."),
        ("Hãy sắp xếp các bước để lấy ngày tháng năm hiện tại và in ra màn hình:", ["import datetime", "now = datetime.datetime.now()", "ngay = now.day", "thang = now.month", "nam = now.year", "print(f'{ngay}/{thang}/{nam}')"], [0, 1, 2, 3, 4, 5], "Quy trình: Import datetime -> Lấy now() -> Tách ngày, tháng, năm -> In định dạng ngày/tháng/năm."),
        ("Hãy sắp xếp các bước để nhấc bút, di chuyển sang vị trí mới và vẽ hình tiếp theo:", ["but_ve.forward(100)", "but_ve.penup()", "but_ve.forward(50)", "but_ve.pendown()", "but_ve.circle(30)"], [0, 1, 2, 3, 4], "Quy trình: Vẽ đoạn 1 -> Nhấc bút penup() -> Dịch chuyển khoảng cách trống -> Hạ bút pendown() -> Vẽ hình tròn.")
    ]

    for item in order_data:
        questions.append({
            "id": q_id,
            "type": "sequence_order",
            "type_name": "Sắp xếp thứ tự logic / đoạn mã",
            "question": item[0],
            "items": item[1],
            "correct_order": item[2],
            "explanation": item[3]
        })
        q_id += 1

    # ----------------------------------------------------
    # DẠNG 6: NỐI QUY TRÌNH / GHÉP CẶP KHÁI NIỆM (10 CÂU: ID 111 -> 120)
    # ----------------------------------------------------
    match_data = [
        ("Hãy ghép cặp giữa Phương thức xử lý chuỗi và Chức năng hoạt động tương ứng:", [("upper()", "Chuyển thành chữ hoa toàn bộ"), ("lower()", "Chuyển thành chữ thường toàn bộ"), ("strip()", "Xóa khoảng trắng thừa ở hai đầu"), ("title()", "Viết hoa chữ cái đầu của mỗi từ")], "Mỗi phương thức chuỗi có một công dụng định dạng chuyên biệt."),
        ("Hãy ghép cặp giữa Tên hàm trong module math và Ý nghĩa toán học:", [("sqrt(x)", "Tính căn bậc hai của số x"), ("pow(a, b)", "Tính lũy thừa a mũ b"), ("factorial(n)", "Tính giai thừa n!"), ("pi", "Hằng số Pi xấp xỉ 3.14159")], "Module math cung cấp đầy đủ các phép toán khoa học chính xác cao."),
        ("Hãy ghép cặp giữa Câu lệnh Turtle và Tác vụ đồ họa tương ứng:", [("forward(d)", "Tiến thẳng d bước"), ("right(angle)", "Xoay phải một góc angle độ"), ("circle(r)", "Vẽ hình tròn bán kính r"), ("penup()", "Nhấc bút không vẽ nét")], "Các câu lệnh cơ bản giúp điều khiển đường đi của bút vẽ Turtle."),
        ("Hãy ghép cặp giữa Phương thức của List và Tác vụ trên danh sách:", [("append(x)", "Thêm phần tử vào cuối danh sách"), ("remove(x)", "Xóa phần tử x đầu tiên tìm thấy"), ("sort()", "Sắp xếp danh sách tăng dần"), ("pop()", "Lấy và xóa phần tử cuối cùng")], "Các phương thức giúp chỉnh sửa và quản trị danh sách linh hoạt."),
        ("Hãy ghép cặp giữa Thư viện Python và Lĩnh vực ứng dụng chính:", [("random", "Sinh số và bốc thăm ngẫu nhiên"), ("math", "Các phép toán nâng cao và lượng giác"), ("datetime", "Xử lý ngày tháng và thời gian thực"), ("turtle", "Lập trình đồ họa vẽ tranh hình học")], "Python có hệ sinh thái thư viện phong phú phục vụ đa dạng nhu cầu."),
        ("Hãy ghép cặp giữa Ký tự điều khiển (Escape Code) và Ý nghĩa hiển thị:", [("\\n", "Xuống dòng mới"), ("\\t", "Thụt lề một khoảng Tab"), ("\\\\", "In ký tự dấu gạch chéo ngược"), ("\\'", "In ký tự dấu nháy đơn trong chuỗi")], "Ký tự escape giúp in các định dạng văn bản đặc biệt."),
        ("Hãy ghép cặp giữa Cú pháp Slicing và Kết quả trích xuất:", [("s[0]", "Lấy ký tự đầu tiên"), ("s[-1]", "Lấy ký tự cuối cùng"), ("s[:3]", "Lấy 3 ký tự đầu tiên"), ("s[::-1]", "Đảo ngược toàn bộ chuỗi")], "Kỹ thuật Slicing là công cụ cắt lọc chuỗi mạnh mẽ nhất trong Python."),
        ("Hãy ghép cặp giữa Cấu trúc lệnh và Mục đích sử dụng trong lập trình:", [("def", "Định nghĩa một hàm tái sử dụng"), ("return", "Trả về giá trị từ hàm"), ("break", "Thoát khỏi vòng lặp ngay lập tức"), ("continue", "Bỏ qua lần lặp hiện tại để sang lần kế tiếp")], "Các từ khóa cốt lõi điều khiển luồng thực thi chương trình."),
        ("Hãy ghép cặp giữa Phương thức kiểm tra chuỗi và Điều kiện trả về True:", [("isdigit()", "Toàn bộ chuỗi là chữ số 0-9"), ("isalpha()", "Toàn bộ chuỗi là chữ cái"), ("isalnum()", "Chuỗi gồm chữ cái và số kết hợp"), ("isspace()", "Toàn bộ chuỗi là khoảng trắng")], "Các hàm tiền tố 'is' giúp kiểm tra tính hợp lệ của dữ liệu người dùng nhập."),
        ("Hãy ghép cặp giữa Câu lệnh màu sắc Turtle và Đối tượng áp dụng:", [("bgcolor('color')", "Đặt màu nền cho toàn bộ cửa sổ"), ("pensize(size)", "Đặt độ dày của nét bút vẽ"), ("begin_fill()", "Bắt đầu vùng tô màu kín"), ("color('c1', 'c2')", "Đặt đồng thời màu viền c1 và màu tô c2")], "Làm chủ các lệnh màu sắc giúp tạo nên những bức tranh đồ họa sinh động.")
    ]

    for item in match_data:
        left_items = [p[0] for p in item[1]]
        right_items = [p[1] for p in item[1]]
        questions.append({
            "id": q_id,
            "type": "matching",
            "type_name": "Nối quy trình / Ghép cặp khái niệm",
            "question": item[0],
            "pairs": [{"left": p[0], "right": p[1]} for p in item[1]],
            "left_items": left_items,
            "right_items": right_items,
            "explanation": item[2]
        })
        q_id += 1

    return questions

def get_10_practical_problems():
    problems = [
        {
            "id": 1,
            "title": "Hàm tính tổng 2 số có kiểm tra đầu vào",
            "description": "Viết hàm tinh_tong(a, b) nhận vào hai số a và b. Trả về tổng của hai số.",
            "starter_code": "def tinh_tong(a, b):\n    # Viết code của em ở đây\n    pass\n\n# Chạy thử kiểm tra\nprint(tinh_tong(15, 25))",
            "solution_code": "def tinh_tong(a, b):\n    return a + b\n\nprint(tinh_tong(15, 25))",
            "test_cases": [
                {"input": "15, 25", "expected_output": "40"},
                {"input": "10.5, 4.5", "expected_output": "15.0"}
            ]
        },
        {
            "id": 2,
            "title": "Hàm kiểm tra số chẵn hay số lẻ",
            "description": "Viết hàm kiem_tra_chan(n) nhận vào một số nguyên n. Trả về True nếu n là số chẵn, ngược lại trả về False.",
            "starter_code": "def kiem_tra_chan(n):\n    # Viết code của em ở đây\n    pass\n\nprint(kiem_tra_chan(8))\nprint(kiem_tra_chan(7))",
            "solution_code": "def kiem_tra_chan(n):\n    return n % 2 == 0\n\nprint(kiem_tra_chan(8))\nprint(kiem_tra_chan(7))",
            "test_cases": [
                {"input": "8", "expected_output": "True"},
                {"input": "7", "expected_output": "False"}
            ]
        },
        {
            "id": 3,
            "title": "Hàm in bảng cửu chương của số n",
            "description": "Viết hàm in_bang_cuu_chuong(n) in ra 10 dòng bảng cửu chương của n từ 1 đến 10 theo mẫu: n x i = ket_qua.",
            "starter_code": "def in_bang_cuu_chuong(n):\n    # Viết code của em ở đây\n    pass\n\nin_bang_cuu_chuong(5)",
            "solution_code": "def in_bang_cuu_chuong(n):\n    for i in range(1, 11):\n        print(f'{n} x {i} = {n * i}')\n\nin_bang_cuu_chuong(5)",
            "test_cases": [
                {"input": "5", "expected_output": "5 x 1 = 5 ... 5 x 10 = 50"}
            ]
        },
        {
            "id": 4,
            "title": "Hàm tính diện tích hình tròn với math.pi",
            "description": "Viết hàm tinh_dien_tich_tron(r) sử dụng hằng số math.pi và trả về diện tích hình tròn có bán kính r.",
            "starter_code": "import math\n\ndef tinh_dien_tich_tron(r):\n    # Viết code của em ở đây\n    pass\n\nprint(round(tinh_dien_tich_tron(5), 2))",
            "solution_code": "import math\n\ndef tinh_dien_tich_tron(r):\n    return math.pi * (r ** 2)\n\nprint(round(tinh_dien_tich_tron(5), 2))",
            "test_cases": [
                {"input": "5", "expected_output": "78.54"}
            ]
        },
        {
            "id": 5,
            "title": "Hàm đảo ngược chuỗi ký tự bằng Slicing",
            "description": "Viết hàm dao_nguoc_chuoi(s) trả về chuỗi đảo ngược của s bằng kỹ thuật Slicing.",
            "starter_code": "def dao_nguoc_chuoi(s):\n    # Viết code của em ở đây\n    pass\n\nprint(dao_nguoc_chuoi('Python'))",
            "solution_code": "def dao_nguoc_chuoi(s):\n    return s[::-1]\n\nprint(dao_nguoc_chuoi('Python'))",
            "test_cases": [
                {"input": "'Python'", "expected_output": "nohtyP"}
            ]
        },
        {
            "id": 6,
            "title": "Hàm tính giai thừa n! của số nguyên",
            "description": "Viết hàm tinh_giai_thua(n) tính n! = 1 * 2 * ... * n. Quy ước 0! = 1.",
            "starter_code": "def tinh_giai_thua(n):\n    # Viết code của em ở đây\n    pass\n\nprint(tinh_giai_thua(5))",
            "solution_code": "def tinh_giai_thua(n):\n    if n == 0 or n == 1:\n        return 1\n    gt = 1\n    for i in range(2, n + 1):\n        gt *= i\n    return gt\n\nprint(tinh_giai_thua(5))",
            "test_cases": [
                {"input": "5", "expected_output": "120"},
                {"input": "0", "expected_output": "1"}
            ]
        },
        {
            "id": 7,
            "title": "Hàm kiểm tra số nguyên tố",
            "description": "Viết hàm kiem_tra_nguyen_to(n) trả về True nếu n là số nguyên tố, ngược lại trả về False.",
            "starter_code": "def kiem_tra_nguyen_to(n):\n    # Viết code của em ở đây\n    pass\n\nprint(kiem_tra_nguyen_to(17))\nprint(kiem_tra_nguyen_to(18))",
            "solution_code": "def kiem_tra_nguyen_to(n):\n    if n < 2:\n        return False\n    for i in range(2, int(n ** 0.5) + 1):\n        if n % i == 0:\n            return False\n    return True\n\nprint(kiem_tra_nguyen_to(17))\nprint(kiem_tra_nguyen_to(18))",
            "test_cases": [
                {"input": "17", "expected_output": "True"},
                {"input": "18", "expected_output": "False"}
            ]
        },
        {
            "id": 8,
            "title": "Hàm tính chu vi và diện tích hình chữ nhật",
            "description": "Viết hàm tinh_hcn(dai, rong) trả về đồng thời 2 giá trị chu vi và diện tích (chu_vi, dien_tich).",
            "starter_code": "def tinh_hcn(dai, rong):\n    # Viết code của em ở đây\n    pass\n\ncv, dt = tinh_hcn(10, 5)\nprint(f'Chu vi: {cv}, Dien tich: {dt}')",
            "solution_code": "def tinh_hcn(dai, rong):\n    return (dai + rong) * 2, dai * rong\n\ncv, dt = tinh_hcn(10, 5)\nprint(f'Chu vi: {cv}, Dien tich: {dt}')",
            "test_cases": [
                {"input": "10, 5", "expected_output": "Chu vi: 30, Dien tich: 50"}
            ]
        },
        {
            "id": 9,
            "title": "Hàm chuyển đổi độ C sang độ F",
            "description": "Viết hàm c_sang_f(c) áp dụng công thức F = (C * 9/5) + 32 và trả về nhiệt độ độ F.",
            "starter_code": "def c_sang_f(c):\n    # Viết code của em ở đây\n    pass\n\nprint(c_sang_f(37))",
            "solution_code": "def c_sang_f(c):\n    return (c * 9 / 5) + 32\n\nprint(c_sang_f(37))",
            "test_cases": [
                {"input": "37", "expected_output": "98.6"},
                {"input": "0", "expected_output": "32.0"}
            ]
        },
        {
            "id": 10,
            "title": "Hàm in danh sách các số chẵn từ 1 đến 100",
            "description": "Viết hàm in_so_chan_1_den_100() sử dụng vòng lặp range(2, 101, 2) in tất cả số chẵn từ 1 đến 100 trên một hàng ngang cách nhau dấu cách.",
            "starter_code": "def in_so_chan_1_den_100():\n    # Viết code của em ở đây\n    pass\n\nin_so_chan_1_den_100()",
            "solution_code": "def in_so_chan_1_den_100():\n    for i in range(2, 101, 2):\n        print(i, end=' ')\n    print()\n\nin_so_chan_1_den_100()",
            "test_cases": [
                {"input": "", "expected_output": "2 4 6 8 ... 98 100"}
            ]
        }
    ]
    return problems

def export_json_and_js(questions, problems):
    data = {
        "title": "Ngân hàng 120 Câu Trắc Nghiệm & 10 Bài Thực Hành Python Nâng Cao",
        "center": "TIN HỌC SAO VIỆT THỦ ĐỨC",
        "total_mcq": len(questions),
        "total_practical": len(problems),
        "questions": questions,
        "practical_problems": problems
    }
    with open(JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"✅ Đã xuất ngân hàng JSON: {JSON_PATH}")

    js_content = f"// Ngân hàng 120 Câu hỏi ôn tập và 10 Bài tập thực hành tự luận Python Nâng Cao\n// Đơn vị: TIN HỌC SAO VIỆT THỦ ĐỨC\nconst QUIZ_DATA = {json.dumps(data, ensure_ascii=False, indent=2)};\n"
    with open(JS_PATH, "w", encoding="utf-8") as f:
        f.write(js_content)
    print(f"✅ Đã xuất dữ liệu JS cho Web: {JS_PATH}")

def build_quiz_docx(questions, problems):
    print("🚀 Bắt đầu tạo tài liệu DOCX Ngân Hàng 120 Câu Hỏi & 10 Bài Thực Hành...")
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(9.5)
    font.color.rgb = BLACK

    # Header Banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "000000")
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("TIN HỌC SAO VIỆT THỦ ĐỨC — TÀI LIỆU ÔN TẬP TOÀN DIỆN")
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("NGÂN HÀNG 120 CÂU TRẮC NGHIỆM & 10 BÀI TẬP THỰC HÀNH TỰ LUẬN")
    run2.font.size = Pt(13.5)
    run2.font.bold = True
    run2.font.color.rgb = WHITE

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Có Đầy Đủ Đáp Án Chuẩn Xác, Giải Thích Suy Luận Logic & Phân Tích Chi Tiết Từng Câu")
    run3.font.size = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = WHITE

    # Divider 1
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(14)
    p_div.paragraph_format.space_after = Pt(4)
    r_d = p_div.add_run("═════ PHẦN 1: KHO 120 CÂU HỎI TRẮC NGHIỆM HỌC THUỘC ÔN TẬP (6 DẠNG) ═════")
    r_d.font.bold = True
    r_d.font.size = Pt(11.5)

    for q in questions:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(2)
        p_q.paragraph_format.keep_with_next = True

        r_badge = p_q.add_run(f"Câu {q['id']} [{q['type_name']}]: ")
        r_badge.font.bold = True
        r_badge.font.size = Pt(10)

        r_text = p_q.add_run(q['question'])
        r_text.font.size = Pt(10)

        # Options formatting based on type
        if q['type'] == 'single_choice':
            labels = ["A", "B", "C", "D"]
            for idx, opt in enumerate(q['options']):
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Inches(0.25)
                p_opt.paragraph_format.space_after = Pt(1)
                p_opt.add_run(f"{labels[idx]}. {opt}")

            ans_idx = q['correct_answer']
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            r_a = p_ans.add_run(f"👉 Đáp án đúng: {labels[ans_idx]}. {q['options'][ans_idx]}")
            r_a.font.bold = True

        elif q['type'] == 'true_false':
            ans_str = "Đúng (True)" if q['correct_answer'] == 0 else "Sai (False)"
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            r_a = p_ans.add_run(f"👉 Đáp án đúng: {ans_str}")
            r_a.font.bold = True

        elif q['type'] == 'multiple_choice':
            labels = ["A", "B", "C", "D"]
            for idx, opt in enumerate(q['options']):
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Inches(0.25)
                p_opt.paragraph_format.space_after = Pt(1)
                p_opt.add_run(f"[  ] {labels[idx]}. {opt}")

            ans_labels = [labels[i] for i in q['correct_answer']]
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            r_a = p_ans.add_run(f"👉 Các đáp án đúng: {', '.join(ans_labels)}")
            r_a.font.bold = True

        elif q['type'] == 'fill_blank':
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            r_a = p_ans.add_run(f"👉 Từ khóa cần điền: '{q['correct_answer']}'")
            r_a.font.bold = True

        elif q['type'] == 'sequence_order':
            p_items = doc.add_paragraph()
            p_items.paragraph_format.left_indent = Inches(0.25)
            p_items.paragraph_format.space_after = Pt(1)
            p_items.add_run("Các bước xáo trộn:\n" + "\n".join([f"• {it}" for it in q['items']]))

            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            ordered_texts = [q['items'][i] for i in q['correct_order']]
            r_a = p_ans.add_run("👉 Thứ tự đúng: " + " ➔ ".join(ordered_texts))
            r_a.font.bold = True

        elif q['type'] == 'matching':
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Inches(0.25)
            p_ans.paragraph_format.space_after = Pt(1)
            r_a = p_ans.add_run("👉 Ghép cặp chính xác:\n" + "\n".join([f"• {p['left']} ── nối với ──▶ {p['right']}" for p in q['pairs']]))
            r_a.font.bold = True

        # Explanation box
        p_exp = doc.add_paragraph()
        p_exp.paragraph_format.left_indent = Inches(0.25)
        p_exp.paragraph_format.space_after = Pt(5)
        r_eh = p_exp.add_run("💡 Chú thích suy luận logic: ")
        r_eh.font.bold = True
        r_eh.font.size = Pt(9)
        r_eb = p_exp.add_run(q['explanation'])
        r_eb.font.size = Pt(9)
        r_eb.font.italic = True

    # Divider 2: 10 Practical Problems
    p_div2 = doc.add_paragraph()
    p_div2.paragraph_format.space_before = Pt(16)
    p_div2.paragraph_format.space_after = Pt(4)
    r_d2 = p_div2.add_run("═════ PHẦN 2: KHO 10 BÀI TẬP THỰC HÀNH TỰ LUẬN ĐỀ THI (40 PHÚT - RANDOM 4 BÀI) ═════")
    r_d2.font.bold = True
    r_d2.font.size = Pt(11.5)

    for p in problems:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(2)
        p_title.paragraph_format.keep_with_next = True
        r_pt = p_title.add_run(f"Bài Thực Hành {p['id']}: {p['title']}")
        r_pt.font.bold = True
        r_pt.font.size = Pt(10.5)

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.2)
        p_desc.paragraph_format.space_after = Pt(2)
        p_desc.add_run(f"• Yêu cầu đề bài: {p['description']}")

        # Solution Code Table
        tbl_code = doc.add_table(rows=1, cols=1)
        tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_code = tbl_code.cell(0, 0)
        set_cell_background(c_code, "FFFFFF")
        set_cell_border(c_code, sz="8", color="000000")
        set_cell_margins(c_code, top=60, bottom=60, left=100, right=100)

        p_c = c_code.paragraphs[0]
        p_c.paragraph_format.space_after = Pt(0)
        r_c = p_c.add_run(p['solution_code'].strip())
        r_c.font.name = "Consolas"
        r_c.font.size = Pt(9.5)

    doc.save(DOCX_PATH)
    print(f"✅ Đã lưu DOCX Ngân hàng câu hỏi: {DOCX_PATH}")

    print("📄 Đang chuyển đổi sang PDF...")
    try:
        convert(DOCX_PATH, PDF_PATH)
        print(f"🎉 Xuất PDF Ngân hàng câu hỏi thành công: {PDF_PATH}")
    except Exception as e:
        print(f"❌ Lỗi khi convert PDF: {e}")

if __name__ == "__main__":
    questions = get_all_120_questions()
    problems = get_10_practical_problems()
    export_json_and_js(questions, problems)
    build_quiz_docx(questions, problems)
